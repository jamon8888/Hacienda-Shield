"""
Hacienda Shield MCP Server v1.0.0
==============================
NER-only mode (no LLM dependency).
  - NER backend: GLiNER (knowledgator/gliner-pii-base-v1.0) via Presidio with SpaCy tokenization
  - High-quality PII-tuned NER via Presidio's GLiNERRecognizer
  - Dependencies managed by UV (pyproject.toml)
  - DRY: single detect() used by both text and docx paths
  - TTL: mappings auto-cleaned after 7 days
  - Unified anonymize_file: routes .docx/.txt/.pdf automatically
  - EU recognizers: 17 patterns (UK, DE, FR, IT, ES, CY, EU-wide)
  - Indexed placeholders: <PERSON_1>, <ORG_1> with reversible mapping
  - Fuzzy entity deduplication, prefix support, boundary cleanup
  - PII-safe responses: mapping and real text never returned to Claude

Tools:
  anonymize_text / anonymize_file / anonymize_docx
  deanonymize_text / deanonymize_docx
  get_mapping / scan_text / list_entities
"""

# ============================================================
# Background model loading  (v1.0 — lifespan + anyio.to_thread)
#
#   mcp.run() blocks plain threading.Thread on Windows (stdin read holds GIL).
#   Fix: use FastMCP lifespan context manager + anyio.to_thread.run_sync()
#   which runs in anyio's managed thread pool — no deadlock.
#
# Server starts accepting MCP connections immediately.
# Tools respond with loading progress until models are ready.
# ============================================================
import subprocess
import sys
import threading
import logging as _boot_log
import time as _time_boot
import os as _os_boot
from contextlib import asynccontextmanager

_boot_log.basicConfig(level=_boot_log.INFO, format="%(asctime)s [Hacienda-Shield] %(message)s", stream=sys.stderr)
_blog = _boot_log.getLogger("hacienda-shield-bootstrap")

_GLINER_MODEL = "knowledgator/gliner-pii-base-v1.0"

# --- Model loading state ---
_engine_ready = threading.Event()
_boot_progress = {"phase": "starting", "message": "Hacienda Shield is starting up...", "pct": 0, "start": None}
_boot_error = None

_HEAVY_PACKAGES = [
    ("presidio_analyzer", "presidio-analyzer>=2.2.355"),
    ("spacy",             "spacy>=3.7.0"),
    ("docx",              "python-docx>=1.1.0"),
    ("cryptography",      "cryptography>=42.0.0"),
    ("numpy",             "numpy>=1.24.0"),
    ("torch",             "torch>=2.0.0"),
    ("gliner",            "gliner>=0.2.7"),
]

# --- Boot benchmark cache for accurate ETA ---
_BENCH_FILE = _os_boot.path.join(_os_boot.path.expanduser("~"), ".hacienda_shield", "boot_benchmark.json")


def _load_boot_benchmark():
    try:
        import json as _j
        with open(_BENCH_FILE, "r") as f:
            return _j.load(f)
    except Exception:
        return {}


def _save_boot_benchmark(timings):
    try:
        import json as _j
        _os_boot.makedirs(_os_boot.path.dirname(_BENCH_FILE), exist_ok=True)
        with open(_BENCH_FILE, "w") as f:
            _j.dump(timings, f)
    except Exception:
        pass


def _all_deps_importable():
    """Quick check via find_spec (no import lock, no deadlock risk)."""
    import importlib.util
    for import_name, _ in _HEAVY_PACKAGES:
        if importlib.util.find_spec(import_name) is None:
            return False
    return True


def _pip_install(packages):
    """Install missing pip packages silently. Returns list of installed specs."""
    missing = []
    for import_name, pip_spec in packages:
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pip_spec)
    if not missing:
        return []
    _blog.info(f"Installing: {missing}")
    log_dir = _os_boot.path.join(_os_boot.path.expanduser("~"), ".hacienda_shield")
    _os_boot.makedirs(log_dir, exist_ok=True)
    log_file = _os_boot.path.join(log_dir, "pip_install.log")
    cflags = subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
    with open(log_file, "a") as lf:
        subprocess.call(
            [sys.executable, "-m", "pip", "install", "--quiet"] + missing,
            stdin=subprocess.DEVNULL, stdout=lf, stderr=lf,
            creationflags=cflags,
        )
    # Verify imports — pip may return non-zero despite success (warnings count as errors)
    still_missing = []
    for import_name, pip_spec in packages:
        try:
            __import__(import_name)
        except ImportError:
            still_missing.append(pip_spec)
    if still_missing:
        raise RuntimeError(f"Failed to install: {still_missing}. Check {log_file}")
    return missing


def _sync_model_load():
    """Heavy init: install deps, load models, init engine. Runs via anyio.to_thread."""
    global _boot_error
    _boot_progress["start"] = _time_boot.monotonic()
    bench = _load_boot_benchmark()
    timings = {}
    MAX_RETRIES = 3

    for attempt in range(1, MAX_RETRIES + 1):
        _boot_error = None
        try:
            return _sync_model_load_inner(bench, timings)
        except OSError as e:
            # Transient errors: DLL lock (WinError 32), file busy, etc.
            if attempt < MAX_RETRIES:
                _blog.warning(f"Attempt {attempt} failed (retrying in 10s): {e}")
                _boot_progress.update(phase="retry", message=f"Transient error, retrying ({attempt}/{MAX_RETRIES})...", pct=5)
                _time_boot.sleep(10)
            else:
                raise


def _sync_model_load_inner(bench, timings):
    """Core init logic. Called by _sync_model_load with retry wrapper."""
    global _boot_error

    try:
        # --- Quick-init vs slow path ---
        if _all_deps_importable():
            _blog.info("Quick-init: all packages present, skipping pip checks")
            _boot_progress.update(phase="models", message="All deps present, loading models...", pct=20)
        else:
            _boot_progress.update(phase="packages", message="Installing missing dependencies...", pct=5)
            _blog.info("Slow path: installing missing packages...")
            installed = _pip_install(_HEAVY_PACKAGES)
            if installed:
                _boot_progress.update(message=f"Installed {len(installed)} packages.", pct=20)
                _blog.info(f"Installed {len(installed)} packages")

        # --- SpaCy tokenizer ---
        eta_spacy = bench.get("spacy_sec", "?")
        _boot_progress.update(phase="models", message=f"Loading SpaCy tokenizer... (prev: {eta_spacy}s)", pct=30)
        _blog.info("Loading SpaCy tokenizer...")
        t0 = _time_boot.monotonic()
        import spacy
        try:
            spacy.load("en_core_web_sm")
        except OSError:
            _blog.info("SpaCy model not found, downloading...")
            cflags = subprocess.CREATE_NO_WINDOW if sys.platform == "win32" else 0
            subprocess.check_call(
                [sys.executable, "-m", "spacy", "download", "en_core_web_sm"],
                stdin=subprocess.DEVNULL, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
                creationflags=cflags,
            )
            spacy.load("en_core_web_sm")
        timings["spacy_sec"] = round(_time_boot.monotonic() - t0, 1)
        _blog.info(f"SpaCy loaded in {timings['spacy_sec']}s")

        # --- GLiNER NER model ---
        gliner_model = _os_boot.environ.get("PII_GLINER_MODEL", _GLINER_MODEL)
        eta_gliner = bench.get("gliner_sec", "?")
        _boot_progress.update(phase="models", message=f"Loading GLiNER ({gliner_model})... (prev: {eta_gliner}s)", pct=55)
        _blog.info(f"Loading GLiNER model: {gliner_model}")
        t0 = _time_boot.monotonic()
        from gliner import GLiNER
        GLiNER.from_pretrained(gliner_model)
        timings["gliner_sec"] = round(_time_boot.monotonic() - t0, 1)
        _blog.info(f"GLiNER loaded in {timings['gliner_sec']}s")

        # --- Initialize PII engine ---
        _boot_progress.update(phase="engine", message="Initializing PII engine...", pct=85)
        _blog.info("Initializing PII engine...")
        for _w in range(30):
            if "engine" in globals():
                break
            _time_boot.sleep(1)
        if "engine" in globals():
            t0 = _time_boot.monotonic()
            engine._ensure_ready(_from_bootstrap=True)
            timings["engine_sec"] = round(_time_boot.monotonic() - t0, 1)
            _blog.info(f"PII engine initialized in {timings['engine_sec']}s")
        else:
            _blog.warning("engine not yet defined, will init on first tool call")

        # Save benchmark for next boot
        timings["total_sec"] = round(_time_boot.monotonic() - _boot_progress["start"], 1)
        _save_boot_benchmark(timings)

        _boot_progress.update(phase="ready", message="Hacienda Shield is ready.", pct=100)
        _blog.info(f"Boot complete in {timings['total_sec']}s — ready for tool calls")

    except Exception as e:
        _boot_error = str(e)
        _boot_progress.update(phase="error", message=f"Failed: {e}", pct=0)
        _blog.error(f"Bootstrap failed: {e}")
    finally:
        _engine_ready.set()


@asynccontextmanager
async def _pii_lifespan(app):
    """FastMCP lifespan: starts model loading in anyio thread pool."""
    import anyio
    async with anyio.create_task_group() as tg:
        async def _run_init():
            await anyio.to_thread.run_sync(_sync_model_load)
        tg.start_soon(_run_init)
        yield {}


# --- Startup: ensure MCP is available (sync, one-time install if needed) ---
try:
    __import__("mcp")
except ImportError:
    _blog.info("MCP not found — installing (one-time)...")
    try:
        _pip_install([("mcp", "mcp[cli]>=1.0.0")])
        __import__("mcp")
    except Exception as _e:
        print(
            f"FATAL: Cannot install MCP package: {_e}\n"
            f"Fix: pip install 'mcp[cli]>=1.0.0'  then restart Hacienda Shield.",
            file=sys.stderr,
        )
        sys.exit(1)
# ============================================================

import json
import os
import re
import time
import uuid
import logging
from pathlib import Path
from collections import defaultdict

from mcp.server.fastmcp import FastMCP
# NOTE: presidio imports are deferred to _ensure_ready() — not yet installed on first run

logging.basicConfig(level=logging.INFO, format="%(asctime)s [Hacienda-Shield] %(message)s", stream=sys.stderr)
log = logging.getLogger("hacienda-shield")

# --- Detailed file logger for diagnostics ---
_flog = logging.getLogger("hacienda-shield-detail")
_flog.setLevel(logging.DEBUG)
_flog.propagate = False  # don't echo to stderr
_flog_handler = None
_FILE_LOG_DISABLED = False

def _ensure_file_log(folder=None):
    """Set up NER log in ~/.hacienda_shield/audit/. Called once per session."""
    global _flog_handler, _FILE_LOG_DISABLED
    if _flog_handler is not None or _FILE_LOG_DISABLED:
        return
    log_dir = Path.home() / ".hacienda_shield" / "audit"
    log_path = log_dir / "ner_debug.log"
    try:
        log_dir.mkdir(parents=True, exist_ok=True)
        _flog_handler = logging.FileHandler(str(log_path), mode="a", encoding="utf-8")
    except OSError as exc:
        _FILE_LOG_DISABLED = True
        log.warning(f"NER log disabled: {exc}")
        return
    _flog_handler.setFormatter(logging.Formatter("%(asctime)s %(message)s"))
    _flog.addHandler(_flog_handler)
    _flog.info(f"===== Hacienda Shield session started =====")
    log.info(f"NER log: {log_path}")

# --- MCP audit logger: logs every tool request/response to prove no PII leaves the machine ---
_audit_log = logging.getLogger("hacienda-shield-audit")
_audit_log.setLevel(logging.DEBUG)
_audit_log.propagate = False
_audit_handler = None
_AUDIT_LOG_DISABLED = False

def _ensure_audit_log():
    """Set up audit logger in ~/.hacienda_shield/audit/. Created once, persists across sessions."""
    global _audit_handler, _AUDIT_LOG_DISABLED
    if _audit_handler is not None or _AUDIT_LOG_DISABLED:
        return
    audit_dir = Path.home() / ".hacienda_shield" / "audit"
    audit_path = audit_dir / "mcp_audit.log"
    try:
        audit_dir.mkdir(parents=True, exist_ok=True)
        _audit_handler = logging.FileHandler(str(audit_path), mode="a", encoding="utf-8")
    except OSError as exc:
        _AUDIT_LOG_DISABLED = True
        log.warning(f"Audit log disabled: {exc}")
        return
    _audit_handler.setFormatter(logging.Formatter("%(asctime)s %(message)s"))
    _audit_log.addHandler(_audit_handler)
    log.info(f"Audit log: {audit_path}")

def _audit_tool(func):
    """Decorator that logs every MCP tool call (args) and response to the audit log."""
    import functools, inspect
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        _ensure_audit_log()
        # Build readable args: match parameter names to values
        sig = inspect.signature(func)
        param_names = list(sig.parameters.keys())
        call_args = {}
        for i, v in enumerate(args):
            if i < len(param_names):
                call_args[param_names[i]] = v
        call_args.update(kwargs)
        # Truncate long text values in the log (the point is to prove they DON'T contain PII)
        safe_args = {}
        for k, v in call_args.items():
            if isinstance(v, str) and len(v) > 500:
                safe_args[k] = v[:200] + f"... [{len(v)} chars total]"
            else:
                safe_args[k] = v
        _audit_log.info(f">>> CALL {func.__name__}({json.dumps(safe_args, ensure_ascii=False)})")
        try:
            result = func(*args, **kwargs)
            # Truncate response too
            if isinstance(result, str) and len(result) > 1000:
                logged_result = result[:500] + f"... [{len(result)} chars total]"
            else:
                logged_result = result
            _audit_log.info(f"<<< RESP {func.__name__} -> {logged_result}")
            return result
        except Exception as exc:
            _audit_log.info(f"<<< ERR  {func.__name__} -> {type(exc).__name__}: {exc}")
            raise
    return wrapper

# ============================================================
# Config
# ============================================================
_DEFAULT_MIN_SCORE = 0.50

def _get_min_score():
    """Read MIN_SCORE from env on every call — allows live config updates."""
    return float(os.environ.get("PII_MIN_SCORE", str(_DEFAULT_MIN_SCORE)))

MIN_SCORE = _get_min_score()  # cached for list_entities display
MAPPING_TTL_DAYS = int(os.environ.get("PII_MAPPING_TTL_DAYS", "7"))

MAPPING_DIR = Path.home() / ".hacienda_shield" / "mappings"
try:
    MAPPING_DIR.mkdir(parents=True, exist_ok=True)
except Exception:
    pass  # will retry in save_mapping; in-memory fallback always works

SUPPORTED_ENTITIES = [
    "PERSON", "ORGANIZATION", "LOCATION", "NRP",
    "EMAIL_ADDRESS", "PHONE_NUMBER", "URL", "IP_ADDRESS",
    "CREDIT_CARD", "IBAN_CODE", "CRYPTO",
    "US_SSN", "US_PASSPORT", "US_DRIVER_LICENSE",
    "UK_NHS", "UK_NIN", "UK_PASSPORT", "UK_CRN", "UK_DRIVING_LICENCE",
    "EU_VAT", "EU_PASSPORT",
    "DE_TAX_ID", "DE_SOCIAL_SECURITY",
    "FR_NIR", "FR_CNI",
    "IT_FISCAL_CODE", "IT_VAT",
    "ES_DNI", "ES_NIE",
    "CY_TIC", "CY_ID_CARD",
    "MEDICAL_LICENSE",
    # Legal (France)
    "FR_SIRET", "FR_SIREN", "FR_RCS", "FR_RG", "FR_NUMERO_PARQUET", "FR_TOQUE",
    # Health (France)
    "FR_RPPS", "FR_ADELI", "FR_FINESS", "FR_CPS_CARD", "FR_NDA",
    # Finance (France)
    "FR_AMF", "FR_LEI", "FR_BBAN", "FR_ACPR",
    # Accounting (France)
    "FR_APE_NAF", "FR_OEC",
    # Semantic — GLiNER zero-shot only (no regex)
    "MEDICAL_CONDITION", "MEDICATION", "MEDICAL_PROCEDURE",
    "LEGAL_FILE_REF", "LEGAL_PRIVILEGE", "FR_BARREAU",
    "FINANCIAL_AMOUNT", "FINANCIAL_INSTRUMENT",
    "INVOICE_REF", "CLIENT_REF",
]

TAG_NAMES = {
    "PERSON": "PERSON", "ORGANIZATION": "ORG", "LOCATION": "LOCATION",
    "NRP": "NRP",
    "EMAIL_ADDRESS": "EMAIL", "PHONE_NUMBER": "PHONE", "URL": "URL",
    "IP_ADDRESS": "IP", "CREDIT_CARD": "CREDIT_CARD", "IBAN_CODE": "IBAN",
    "CRYPTO": "CRYPTO",
    "US_SSN": "US_SSN", "US_PASSPORT": "US_PASSPORT", "US_DRIVER_LICENSE": "US_DL",
    "UK_NHS": "UK_NHS", "UK_NIN": "UK_NIN", "UK_PASSPORT": "UK_PASSPORT",
    "UK_CRN": "UK_CRN", "UK_DRIVING_LICENCE": "UK_DL",
    "EU_VAT": "EU_VAT", "EU_PASSPORT": "EU_PASSPORT",
    "DE_TAX_ID": "DE_TAX", "DE_SOCIAL_SECURITY": "DE_SSN",
    "FR_NIR": "FR_NIR", "FR_CNI": "FR_CNI",
    "IT_FISCAL_CODE": "IT_CF", "IT_VAT": "IT_VAT",
    "ES_DNI": "ES_DNI", "ES_NIE": "ES_NIE",
    "CY_TIC": "CY_TIC", "CY_ID_CARD": "CY_ID",
    "MEDICAL_LICENSE": "MED_LIC",
    "FR_SIRET": "FR_SIRET", "FR_SIREN": "FR_SIREN", "FR_RCS": "FR_RCS",
    "FR_RG": "FR_RG", "FR_NUMERO_PARQUET": "FR_PARQUET", "FR_TOQUE": "FR_TOQUE",
    "FR_RPPS": "FR_RPPS", "FR_ADELI": "FR_ADELI", "FR_FINESS": "FR_FINESS",
    "FR_CPS_CARD": "FR_CPS", "FR_NDA": "FR_NDA",
    "FR_AMF": "FR_AMF", "FR_LEI": "LEI", "FR_BBAN": "FR_BBAN", "FR_ACPR": "FR_ACPR",
    "FR_APE_NAF": "FR_NAF", "FR_OEC": "FR_OEC",
    "MEDICAL_CONDITION": "MED_COND", "MEDICATION": "MEDICATION",
    "MEDICAL_PROCEDURE": "MED_PROC",
    "LEGAL_FILE_REF": "LEGAL_REF", "LEGAL_PRIVILEGE": "LEGAL_PRIV",
    "FR_BARREAU": "FR_BARREAU",
    "FINANCIAL_AMOUNT": "FIN_AMOUNT", "FINANCIAL_INSTRUMENT": "FIN_INST",
    "INVOICE_REF": "INVOICE", "CLIENT_REF": "CLIENT_REF",
}


# ============================================================
# Mapping persistence + TTL
# ============================================================
_in_memory_mappings = {}

def save_mapping(session_id, mapping, metadata=None):
    data = {"session_id": session_id, "mapping": mapping, "metadata": metadata or {}, "timestamp": time.time()}
    # Always keep in memory FIRST — this never fails
    _in_memory_mappings[session_id] = data
    # Try to persist to disk (optional — in-memory is the primary store)
    disk_path = None
    try:
        MAPPING_DIR.mkdir(parents=True, exist_ok=True)
        path = MAPPING_DIR / f"{session_id}.json"
        tmp = path.with_suffix(".tmp")
        tmp.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
        tmp.replace(path)
        disk_path = str(path)
    except Exception as e:
        log.warning(f"save_mapping disk write failed (in-memory OK): {e}")
    return disk_path or f"memory://{session_id}"


def load_mapping(session_id):
    # Try disk first
    try:
        path = MAPPING_DIR / f"{session_id}.json"
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8")).get("mapping", {})
    except Exception as e:
        log.warning(f"load_mapping disk read failed: {e}")
    # Fallback to in-memory
    if session_id in _in_memory_mappings:
        return _in_memory_mappings[session_id].get("mapping", {})
    return {}


def cleanup_old_mappings():
    """Delete mappings older than TTL."""
    cutoff = time.time() - (MAPPING_TTL_DAYS * 86400)
    removed = 0
    for f in MAPPING_DIR.glob("*.json"):
        try:
            if f.stat().st_mtime < cutoff:
                f.unlink()
                removed += 1
        except Exception:
            pass
    if removed:
        log.info(f"Cleaned up {removed} expired mappings (>{MAPPING_TTL_DAYS} days)")


def _save_review_to_disk(session_id, review_data):
    """Persist review data to disk so other server processes can access it."""
    try:
        path = MAPPING_DIR / f"review_{session_id}.json"
        tmp = path.with_suffix(".tmp")
        tmp.write_text(json.dumps(review_data, ensure_ascii=False), encoding="utf-8")
        tmp.replace(path)
    except Exception as e:
        log.warning(f"_save_review_to_disk failed (in-memory OK): {e}")


def _load_review_from_disk(session_id):
    """Load review data from disk (cross-process fallback)."""
    try:
        path = MAPPING_DIR / f"review_{session_id}.json"
        if path.exists():
            data = json.loads(path.read_text(encoding="utf-8"))
            # Cache in memory for fast subsequent access
            _in_memory_mappings[f"review:{session_id}"] = data
            return data
    except Exception as e:
        log.warning(f"_load_review_from_disk failed: {e}")
    return None


def _get_review(session_id):
    """Get review data: memory first, then disk."""
    review_key = f"review:{session_id}"
    if review_key in _in_memory_mappings:
        return _in_memory_mappings[review_key]
    return _load_review_from_disk(session_id)


# ============================================================
# PIIEngine
# ============================================================
class PIIEngine:
    _instance = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance

    GLINER_MODEL_NAME = os.environ.get("PII_GLINER_MODEL", _GLINER_MODEL)

    def _ensure_ready(self, _from_bootstrap=False):
        """Initialize the PII engine. Called from background thread after models are loaded,
        or lazily on first tool call.

        Uses GLiNER PII-tuned NER for high-quality entity recognition.
        SpaCy handles tokenization, GLiNER handles NER via Presidio's GLiNERRecognizer.
        Falls back to SpaCy-only if GLiNER is unavailable.
        """
        if self._initialized:
            return

        # If called from a tool (not from background loader), wait for models to load first.
        if not _from_bootstrap:
            if not _engine_ready.is_set():
                log.info("Waiting for background model load to complete...")
                if not _engine_ready.wait(timeout=600):
                    raise RuntimeError(
                        "Model load timed out after 10 minutes. "
                        "Check internet connection and try restarting."
                    )
            if _boot_error:
                log.warning(f"Background load had errors: {_boot_error}")

        log.info("Initializing PII Engine v1.0.0...")

        # --- SpaCy NLP engine (tokenization only) ---
        from presidio_analyzer.nlp_engine import NlpEngineProvider
        try:
            nlp_engine = NlpEngineProvider(nlp_configuration={
                "nlp_engine_name": "spacy",
                "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
            }).create_engine()
        except Exception as e:
            raise RuntimeError(f"SpaCy engine failed: {e}. Ensure spacy and en_core_web_sm are installed.") from e

        # --- Build registry with built-in recognizers ---
        from presidio_analyzer import AnalyzerEngine, RecognizerRegistry
        registry = RecognizerRegistry()
        registry.load_predefined_recognizers()
        # Remove noisy recognizers:
        # - DateRecognizer: too aggressive on legal docs (catches "30 days", section numbers, years)
        # - SpacyRecognizer: flat score=0.85 on everything, generates 90%+ false positives.
        #   GLiNER handles NER with meaningful confidence scores. SpaCy stays as tokenizer only.
        _remove_recognizers = {"DateRecognizer", "SpacyRecognizer"}
        registry.recognizers = [r for r in registry.recognizers if type(r).__name__ not in _remove_recognizers]
        log.info(f"Loaded {len(registry.recognizers)} predefined recognizers (DateRecognizer + SpacyRecognizer removed)")

        # --- Try GLiNER (zero-shot NER), fallback to SpaCy-only ---
        backend_used = "spacy (en_core_web_sm) [FALLBACK]"
        try:
            from presidio_analyzer.predefined_recognizers import GLiNERRecognizer
            gliner_recognizer = GLiNERRecognizer(
                model_name=self.GLINER_MODEL_NAME,
                entity_mapping={
                    # Named entities
                    "person": "PERSON",
                    "name": "PERSON",
                    "company": "ORGANIZATION",
                    "organization": "ORGANIZATION",
                    "location": "LOCATION",
                    "address": "LOCATION",
                    "nationality": "NRP",
                    # Contact info
                    "phone number": "PHONE_NUMBER",
                    "mobile phone number": "PHONE_NUMBER",
                    "landline phone number": "PHONE_NUMBER",
                    "fax number": "PHONE_NUMBER",
                    "email": "EMAIL_ADDRESS",
                    "email address": "EMAIL_ADDRESS",
                    # Financial
                    "credit card number": "CREDIT_CARD",
                    "iban": "IBAN_CODE",
                    "bank account number": "IBAN_CODE",
                    # IDs
                    "social security number": "US_SSN",
                    "passport number": "US_PASSPORT",
                    "driver's license number": "US_DRIVER_LICENSE",
                    "tax identification number": "DE_TAX_ID",
                    "national id number": "CY_ID_CARD",
                    "identity card number": "CY_ID_CARD",
                    # Digital
                    "ip address": "IP_ADDRESS",
                    "url": "URL",
                    "username": "URL",
                    # Medical
                    "medical condition": "MEDICAL_CONDITION",
                    "health insurance id number": "UK_NHS",
                    # Date
                    "date of birth": "DATE_TIME",
                    # Health — semantic detection (no regex possible for these)
                    "medical diagnosis": "MEDICAL_CONDITION",
                    "health condition": "MEDICAL_CONDITION",
                    "disease": "MEDICAL_CONDITION",
                    "prescribed medication": "MEDICATION",
                    "drug name": "MEDICATION",
                    "medical procedure": "MEDICAL_PROCEDURE",
                    "patient identifier": "FR_NDA",
                    "RPPS number": "FR_RPPS",
                    "ADELI number": "FR_ADELI",
                    "FINESS number": "FR_FINESS",
                    # Legal — privilege is semantic
                    "case file number": "LEGAL_FILE_REF",
                    "court reference": "LEGAL_FILE_REF",
                    "numéro de rôle": "FR_RG",
                    "client instruction": "LEGAL_PRIVILEGE",
                    "legal opinion": "LEGAL_PRIVILEGE",
                    "bar registration number": "FR_BARREAU",
                    # Finance — ACPR and OEC have no reliable regex
                    "account balance": "FINANCIAL_AMOUNT",
                    "salary amount": "FINANCIAL_AMOUNT",
                    "investment portfolio": "FINANCIAL_INSTRUMENT",
                    "AMF registration number": "FR_AMF",
                    "LEI code": "FR_LEI",
                    "SIRET number": "FR_SIRET",
                    "SIREN number": "FR_SIREN",
                    "ACPR registration number": "FR_ACPR",
                    "bank routing code": "FR_BBAN",
                    # Accounting — OEC has no reliable regex
                    "invoice number": "INVOICE_REF",
                    "client file number": "CLIENT_REF",
                    "NAF code": "FR_APE_NAF",
                    "accounting professional number": "FR_OEC",
                },
                flat_ner=False,
                multi_label=True,
                map_location="cpu",
            )
            registry.add_recognizer(gliner_recognizer)
            backend_used = f"gliner ({self.GLINER_MODEL_NAME})"
            log.info(f"GLiNER recognizer loaded: {backend_used}")
        except Exception as e:
            log.warning(f"GLiNER failed: {e}")
            log.warning("Using SpaCy-only NER (reduced quality)")

        self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine, registry=registry)
        self._backend = backend_used

        # EU recognizers
        try:
            import importlib.util
            eu_path = Path(__file__).parent / "eu_recognizers.py"
            if eu_path.exists():
                spec = importlib.util.spec_from_file_location("eu_recognizers", eu_path)
                mod = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(mod)
                count = mod.register_eu_recognizers(self.analyzer)
                log.info(f"EU recognizers: {count} registered")
        except Exception as e:
            log.warning(f"EU recognizers failed: {e}")

        self._initialized = True

        cleanup_old_mappings()
        log.info(f"PII Engine ready ({backend_used})")

    def __init__(self):
        # Lightweight — just mark as not initialized.
        # Heavy work is deferred to _ensure_ready() on first tool call.
        pass

    # --- Core detect ---
    def _deduplicate(self, results):
        if not results:
            return []
        s = sorted(results, key=lambda r: (r.start, -r.score))
        d = [s[0]]
        for r in s[1:]:
            if r.start >= d[-1].end:
                d.append(r)
            elif r.score > d[-1].score:
                d[-1] = r
        return d

    # Entity types that are proper-noun based (names, orgs, locations)
    _NAMED_ENTITY_TYPES = {"PERSON", "ORGANIZATION", "LOCATION", "NRP"}

    @staticmethod
    def _snap_word_boundaries(text, entities):
        """Pass 1: Snap entity boundaries to word edges.

        If NER cut a word in the middle, extend to complete that word.
        Never adds new words — only finishes the word already partially captured.
        Never snaps across newlines (paragraph boundaries).
        Also trims trailing/leading punctuation.
        Drops entities that are too short (<=2 chars) after cleanup.
        """
        tlen = len(text)
        _split_buf = []  # collect newline-split entities here (not inside loop)
        for e in entities:
            start, end = e["start"], e["end"]

            # Snap RIGHT: if boundary is mid-word, complete the word
            # Stop at newlines — don't merge across paragraphs
            if end < tlen and end > 0 and text[end].isalnum() and text[end - 1].isalnum():
                while end < tlen and text[end].isalnum() and text[end] != '\n':
                    end += 1

            # Snap LEFT: if boundary is mid-word, complete the word
            if start > 0 and start < end and text[start].isalnum() and text[start - 1].isalnum():
                while start > 0 and text[start - 1].isalnum() and text[start - 1] != '\n':
                    start -= 1

            # Trim trailing/leading punctuation, whitespace, newlines
            while end > start and text[end - 1] in '.,;:)]\'" \t\n\r':
                end -= 1
            while start < end and text[start] in '([\'" \t\n\r#/':
                start += 1

            entity_text = text[start:end].strip()
            if len(entity_text) <= 2:
                # Too short to be meaningful PII (e.g. "S", "St")
                e["_drop"] = True
                log.info(f"Boundary drop (too short): '{entity_text}' "
                         f"(type={e.get('type', '?')})")
            elif '\n' in entity_text:
                # Entity spans multiple lines — split into separate entities per line
                # NOTE: collect into _split_buf to avoid mutating list during iteration
                e["_drop"] = True
                lines = entity_text.split('\n')
                search_from = start
                for line in lines:
                    stripped = line.strip()
                    if len(stripped) > 2:
                        line_start = text.find(stripped, search_from)
                        if line_start == -1:
                            continue  # skip if not found (safety)
                        _split_buf.append({
                            "start": line_start,
                            "end": line_start + len(stripped),
                            "text": stripped,
                            "type": e.get("type", ""),
                            "score": e.get("score", 0),
                        })
                        search_from = line_start + len(stripped)
                    else:
                        # Advance search position past this line
                        pos = text.find(line, search_from)
                        if pos != -1:
                            search_from = pos + len(line)
                log.info(f"Boundary split (across newline): '{entity_text[:40]}' → "
                         f"{[l.strip() for l in lines if len(l.strip()) > 2]}")
            elif start < end:
                e["start"] = start
                e["end"] = end
                e["text"] = entity_text
            else:
                e["_drop"] = True

        # Append split entities collected during the loop (avoids list mutation during iteration)
        entities.extend(_split_buf)
        return [e for e in entities if not e.get("_drop")]

    # Common legal/contract terms that NER frequently misclassifies as PII.
    # Case-insensitive match against normalized entity text.
    _LEGAL_STOPLIST = {
        # ── Contract parties / roles ──
        "contractor", "subcontractor", "client", "customer", "vendor",
        "supplier", "distributor", "franchisor", "franchisee",
        "licensor", "licensee", "employer", "employee", "consultant",
        "agent", "principal", "assignee", "assignor",
        "guarantor", "beneficiary", "trustee", "grantor", "grantee",
        "lessee", "lessor", "tenant", "landlord", "borrower", "lender",
        "buyer", "seller", "partner", "shareholder", "director",
        "officer", "secretary", "treasurer", "representative",
        "obligor", "obligee", "indemnitor", "indemnitee",
        "party", "parties", "counterparty",
        # ── Job titles / corporate roles (NER → PERSON) ──
        "chairman", "chairwoman", "chairperson", "president",
        "vice president", "manager", "supervisor", "administrator",
        "coordinator", "counsel", "attorney", "auditor", "comptroller",
        "commissioner", "mediator", "arbitrator", "notary",
        "general counsel", "key employee", "key employees",
        "ceo", "cfo", "cto", "coo", "cmo", "cio", "cpo",
        # ── Document / legal structural terms ──
        "order", "agreement", "contract", "amendment", "addendum",
        "exhibit", "schedule", "appendix", "annex", "section",
        "article", "clause", "paragraph", "recital", "preamble",
        "purchase order", "statement of work", "scope of work",
        "whereas", "herein", "thereof", "therein", "hereby",
        "definitions", "interpretation", "counterparts", "announcements",
        "variation", "assignment", "notices", "costs",
        # ── M&A / SPA / corporate transaction terms ──
        "shares", "share", "sale", "completion", "conditions",
        "warranties", "warranty", "representations", "covenants",
        "obligations", "undertakings", "indemnities", "limitations",
        "transaction", "acquisition", "disposal", "transfer",
        "consideration", "purchase price", "closing", "escrow",
        "due diligence", "disclosure", "material adverse",
        "pre-completion", "post-completion", "longstop",
        "lien", "encumbrance", "pledge", "charge", "mortgage",
        "de minimis", "de minimis amount", "basket", "cap",
        "tax", "taxation", "tax covenant", "tax deed",
        "hmrc", "customs", "revenue",
        # ── Legal concepts (capitalized in contracts → NER false positives) ──
        "effective date", "termination date", "commencement date",
        "governing law", "force majeure", "confidential information",
        "intellectual property", "indemnification", "arbitration",
        "term", "territory", "termination", "jurisdiction",
        "liability", "negligence", "damages",
        "breach", "remedy", "waiver", "severability",
        "claim", "claims", "dispute", "proceedings", "litigation",
        "consent", "approval", "authority", "resolution",
        # ── Generic business / corporate terms ──
        "company", "corporation", "entity", "firm", "business",
        "affiliate", "subsidiary", "parent", "division", "branch",
        "enterprise", "venture", "consortium", "syndicate",
        "board", "committee", "department", "office",
        "body corporate", "government", "association", "partnership",
        # ── Generic nouns NER misclassifies as PERSON ──
        "person", "individual", "persons", "individuals",
        "actor", "actors", "creator", "creators", "model", "models",
        "influencer", "influencers", "talent", "talents",
        "candidate", "applicant", "recipient", "subscriber",
        "member", "participant", "attendee", "user", "owner",
        "author", "editor", "contributor", "reviewer", "approver",
        "sender", "receiver", "holder", "bearer", "maker",
        "performer", "speaker", "presenter", "moderator",
        "witness", "signatory", "undersigned",
        "purchase", "invoice", "payment", "delivery", "shipment",
        "name", "practice", "relevant person",
        # ── Short ambiguous words (SpaCy/GLiNER false positives) ──
        "will", "may", "case", "show", "set", "lead", "head",
        "share", "note", "record", "draft", "release", "notice",
        # ── Abbreviations ──
        "cta", "nda", "sow", "msa", "sla", "roi", "kpi",
        "llc", "ltd", "inc", "corp", "plc", "gmbh", "sarl", "llp",
        "usd", "eur", "gbp", "jpy", "cny",
        # ── Software / product / brand names (not PII) ──
        "adobe", "adobe premiere", "adobe premiere pro", "adobe after effects",
        "final cut", "final cut pro", "davinci resolve",
        "photoshop", "illustrator", "figma", "canva",
        "microsoft", "google", "apple", "amazon", "meta",
        # ── Cyrillic homoglyphs (С = Cyrillic Es looks like Latin C) ──
        "\u0441lient", "сlient",  # Cyrillic С + lient
    }

    @staticmethod
    def _filter_false_positives(entities):
        """Pass 2: Drop false positives using cross-entity context and stop-list.

        Rules:
        1. Stop-list: known legal/contract terms → always drop.
        2. Single lowercase word + named entity type → drop (not a proper noun).
        3. If same text appears in another entity with higher score → confirms both.
        """
        # Collect all confirmed high-score entity texts for cross-reference
        confirmed_texts = set()
        for e in entities:
            if e.get("score", 0) >= 0.6:
                confirmed_texts.add(e["text"].lower())
                for word in e["text"].split():
                    confirmed_texts.add(word.lower())

        cleaned = []
        for e in entities:
            txt = e["text"]
            etype = e.get("type", "")
            words = txt.split()
            norm_txt = txt.lower().strip()

            # Rule 0: Stop-list — known non-PII terms
            # Also check with Cyrillic homoglyph normalization (С→C, А→A, etc.)
            _CYRILLIC_TO_LATIN = str.maketrans('СсАаЕеОоРрХхВвМмТтНн', 'CcAaEeOoPpXxBbMmTtHh')
            norm_latin = norm_txt.translate(_CYRILLIC_TO_LATIN)
            # Strip leading articles before stoplist check
            _ARTICLES = ("the ", "a ", "an ")
            stripped = norm_txt
            for art in _ARTICLES:
                if stripped.startswith(art):
                    stripped = stripped[len(art):]
                    break
            stripped_latin = norm_latin
            for art in _ARTICLES:
                if stripped_latin.startswith(art):
                    stripped_latin = stripped_latin[len(art):]
                    break
            _sl = PIIEngine._LEGAL_STOPLIST
            if (norm_txt in _sl or norm_latin in _sl
                    or stripped in _sl or stripped_latin in _sl):
                log.info(f"FP drop (stop-list): '{txt}' (type={etype})")
                continue
            # Also drop if ALL words are function words or in stoplist
            _skip_words = {"the", "a", "an", "of", "and", "or", "for", "in", "to", "by",
                           "on", "at", "is", "it", "as", "if", "so", "no", "not", "its",
                           "this", "that", "with", "from", "but", "all", "any", "each",
                           "such", "than", "into", "upon", "per", "via", "re", "vs"}
            meaningful = [w for w in norm_txt.split() if w not in _skip_words]
            if not meaningful:
                # Entity is entirely function words (e.g. "the", "a", "an") — never PII
                log.info(f"FP drop (function words only): '{txt}' (type={etype})")
                continue
            if all(w in _sl for w in meaningful):
                log.info(f"FP drop (all words in stop-list): '{txt}' (type={etype})")
                continue

            # Rule 1: Single lowercase word + named entity type → likely false positive
            if len(words) == 1 and etype in PIIEngine._NAMED_ENTITY_TYPES:
                if txt[0].islower():
                    if txt.lower() not in confirmed_texts:
                        log.info(f"FP drop (single lowercase word): '{txt}' "
                                 f"(type={etype}, score={e.get('score', '?')})")
                        continue

            # Rule 2: Pattern-based FP — short text matched by document-type recognizers
            # EU_VAT/UK_DRIVING_LICENCE/DE_SOCIAL_SECURITY often match "Tax", "Taxation" etc.
            _NOISY_PATTERN_TYPES = {
                "DE_SOCIAL_SECURITY", "EU_VAT", "UK_DRIVING_LICENCE",
                "MEDICAL_LICENSE", "NRP",
            }
            if etype in _NOISY_PATTERN_TYPES and norm_txt in _sl:
                log.info(f"FP drop (noisy pattern + stoplist): '{txt}' (type={etype})")
                continue

            # Rule 3: "Schedule N", "Clause N", "Section N" — structural references, not PII
            import re as _re
            if _re.match(r'^(schedule|clause|section|article|appendix|annex|exhibit|part|recital)\s+\d', norm_txt):
                log.info(f"FP drop (structural reference): '{txt}' (type={etype})")
                continue

            # Rule 4: ALL-CAPS single word ≤10 chars that's in stoplist — section heading
            if txt.isupper() and len(txt) <= 12 and norm_txt in _sl:
                log.info(f"FP drop (all-caps heading): '{txt}' (type={etype})")
                continue

            cleaned.append(e)

        # Rule 5: Frequency filter — if same text appears as entity >8 times,
        # it's almost certainly a structural/legal term, not a real PII name.
        # Real names rarely appear 10+ times; "Company"/"Buyer"/"Seller" do.
        from collections import Counter
        text_counts = Counter(e["text"].lower().strip() for e in cleaned
                              if e.get("type") in PIIEngine._NAMED_ENTITY_TYPES)
        high_freq = {t for t, c in text_counts.items() if c > 8}
        if high_freq:
            before = len(cleaned)
            cleaned = [e for e in cleaned
                       if e["text"].lower().strip() not in high_freq
                       or e.get("type") not in PIIEngine._NAMED_ENTITY_TYPES]
            if len(cleaned) < before:
                log.info(f"FP drop (frequency >8): removed {before - len(cleaned)} entities, "
                         f"terms: {high_freq}")

        return cleaned

    @classmethod
    def _clean_boundaries(cls, text, entities):
        """Two-pass boundary cleanup: snap words, then filter false positives."""
        entities = cls._snap_word_boundaries(text, entities)
        entities = cls._filter_false_positives(entities)
        return entities

    def _analyze_chunked(self, text, language="en", chunk_size=4000, overlap=250):
        """Run analyzer on text in chunks to avoid GLiNER timeout on long texts.
        Chunks overlap to avoid splitting entities at boundaries.
        chunk_size=4000 keeps total time under MCP 60s timeout for ~20-page docs."""
        if len(text) <= chunk_size:
            return self.analyzer.analyze(text=text, entities=SUPPORTED_ENTITIES, language=language)

        all_results = []
        start = 0
        chunk_num = 0
        total_chunks = (len(text) + chunk_size - 1) // chunk_size
        _flog.info(f"  Chunked analysis: ~{total_chunks} chunks, chunk_size={chunk_size}, overlap={overlap}")
        t_chunk_start = time.time()
        while start < len(text):
            end = min(start + chunk_size, len(text))
            # Try to break at whitespace
            if end < len(text):
                ws = text.rfind(' ', start + chunk_size - overlap, end)
                if ws > start:
                    end = ws + 1
            chunk = text[start:end]
            t0 = time.time()
            chunk_results = self.analyzer.analyze(text=chunk, entities=SUPPORTED_ENTITIES, language=language)
            chunk_num += 1
            _flog.info(f"    Chunk {chunk_num}/{total_chunks}: [{start}:{end}] ({end-start} chars) → {len(chunk_results)} detections in {time.time()-t0:.1f}s")
            # Adjust offsets to full text positions
            for r in chunk_results:
                r.start += start
                r.end += start
                all_results.append(r)
            start = end - overlap if end < len(text) else len(text)
        _flog.info(f"  Chunked analysis done: {chunk_num} chunks, {len(all_results)} raw detections in {time.time()-t_chunk_start:.1f}s")

        # Deduplicate overlapping detections (same span)
        seen = set()
        unique = []
        for r in sorted(all_results, key=lambda x: (x.start, -x.score)):
            key = (r.start, r.end, r.entity_type)
            if key not in seen:
                seen.add(key)
                unique.append(r)
        return unique

    def detect(self, text, language="en"):
        """Detect PII using NER. All entities above MIN_SCORE are confirmed."""
        self._ensure_ready()
        min_score = _get_min_score()

        _flog.info(f"--- detect() start | text length={len(text)} | min_score={min_score} | backend={self._backend} ---")

        results = self._analyze_chunked(text, language)
        _flog.info(f"  Raw analyzer results: {len(results)} detections")

        # Log ALL raw results with recognizer info
        for r in results:
            rname = getattr(r, 'recognition_metadata', {}).get('recognizer_name', 'unknown')
            et = text[r.start:r.end]
            _flog.info(f"  RAW: [{rname}] {r.entity_type} score={r.score:.3f} [{r.start}:{r.end}] \"{et}\"")

        results = self._deduplicate(results)
        _flog.info(f"  After dedup: {len(results)} detections")

        entities = []
        skipped_low_score = []

        for r in results:
            et = text[r.start:r.end]
            etype = r.entity_type
            rname = getattr(r, 'recognition_metadata', {}).get('recognizer_name', 'unknown')

            if r.score < min_score:
                skipped_low_score.append(f"[{rname}] {etype}({r.score:.3f})=\"{et}\"")
                continue

            entry = {
                "text": et, "type": etype, "start": r.start, "end": r.end,
                "score": round(r.score, 3), "verified": True, "reason": "NER",
                "_recognizer": rname,
            }
            entities.append(entry)

        if skipped_low_score:
            _flog.info(f"  Skipped (score < {min_score}): {len(skipped_low_score)}")
            for s in skipped_low_score:
                _flog.info(f"    SKIP: {s}")

        # Log confirmed entities with recognizer
        _flog.info(f"  Confirmed entities: {len(entities)}")
        for e in entities:
            _flog.info(f"    OK: [{e['_recognizer']}] {e['type']} score={e['score']} [{e['start']}:{e['end']}] \"{e['text']}\"")

        # Stderr summary
        post_stats = {}
        for e in entities:
            rn = e.get("_recognizer", "unknown")
            post_stats[rn] = post_stats.get(rn, 0) + 1
        log.info(f"Detect: {len(entities)} confirmed, {len(skipped_low_score)} skipped (min_score={min_score}) — by recognizer: {post_stats}")

        for e in entities:
            e.pop("_recognizer", None)

        before_cleanup = len(entities)
        entities = self._clean_boundaries(text, entities)
        if len(entities) != before_cleanup:
            _flog.info(f"  Boundary cleanup: {before_cleanup} → {len(entities)} entities")
            for e in entities:
                _flog.info(f"    FINAL: {e['type']} [{e['start']}:{e['end']}] \"{e['text']}\"")

        _flog.info(f"--- detect() done | {len(entities)} entities ---")
        return entities

    # --- Fuzzy entity deduplication helpers ---
    @staticmethod
    def _normalize(text):
        """Normalize entity text for dedup: lowercase, strip punctuation, collapse spaces."""
        return re.sub(r'\s+', ' ', text.lower().strip().rstrip('.,;:'))

    def _get_or_create_placeholder(self, etype, text, type_counters, seen_exact, seen_family, mapping, prefix=""):
        """Get existing placeholder for this exact entity text, or create new one.

        Exact-match dedup: "Acme" always maps to the same placeholder.
        Family grouping: "Acme" → <ORG_1>, "Acme Corp." → <ORG_1a>,
        "Acme Corporation" → <ORG_1b>. Each has its own placeholder
        with its own reverse mapping, so deanonymization restores exact text.

        Args:
            etype: entity type string
            text: raw entity text as found in document
            type_counters: dict tracking next family index per type
            seen_exact: dict of (type, exact_normalized_text) → placeholder
            seen_family: dict of (type, short_norm) → (family_number, variant_counter)
            mapping: dict of placeholder → raw text
            prefix: optional prefix for multi-file workflows
        Returns:
            placeholder string
        """
        norm = self._normalize(text)
        exact_key = (etype, norm)

        # 1. Exact match — same text seen before, reuse placeholder
        if exact_key in seen_exact:
            return seen_exact[exact_key]

        tag = TAG_NAMES.get(etype, etype)

        # 2. Check if this belongs to an existing family (substring match)
        family_key = None
        if len(norm) >= 4:
            for (ft, fn), (fnum, _) in seen_family.items():
                if ft != etype:
                    continue
                if len(fn) >= 4 and (norm in fn or fn in norm):
                    family_key = (ft, fn)
                    break

        if family_key:
            # Add as variant to existing family
            fnum, vcounter = seen_family[family_key]
            vcounter += 1
            seen_family[family_key] = (fnum, vcounter)
            # variant suffix: a, b, c, ...
            suffix = chr(ord('a') + vcounter - 1) if vcounter <= 26 else str(vcounter)
            if prefix:
                placeholder = f"<{prefix}_{tag}_{fnum}{suffix}>"
            else:
                placeholder = f"<{tag}_{fnum}{suffix}>"
        else:
            # New family
            type_counters[etype] += 1
            fnum = type_counters[etype]
            if prefix:
                placeholder = f"<{prefix}_{tag}_{fnum}>"
            else:
                placeholder = f"<{tag}_{fnum}>"
            # Register this as the family root (shortest/first form)
            seen_family[(etype, norm)] = (fnum, 0)

        seen_exact[exact_key] = placeholder
        mapping[placeholder] = text  # exact text for this specific placeholder
        log.info(f"Placeholder: '{text}' → {placeholder}")
        return placeholder

    # --- Assign indexed placeholders (shared logic) ---
    def _assign_placeholders(self, confirmed_entities, prefix=""):
        """Assign indexed placeholders preserving exact entity forms. Returns mapping dict."""
        type_counters = defaultdict(int)
        seen_exact = {}   # (type, exact_normalized_text) → placeholder
        seen_family = {}  # (type, normalized_text) → (family_number, variant_counter)
        mapping = {}      # placeholder → exact raw text

        for e in sorted(confirmed_entities, key=lambda x: x["start"]):
            e["placeholder"] = self._get_or_create_placeholder(
                e["type"], e["text"], type_counters, seen_exact, seen_family, mapping, prefix
            )

        return mapping

    # --- Apply user overrides (HITL review) ---
    def _apply_overrides(self, confirmed, text, overrides_json):
        """Apply user corrections from HITL review: remove false positives, add missed entities.
        Removes/adds apply to ALL occurrences of the same text+type, not just the clicked one."""
        try:
            overrides = json.loads(overrides_json) if isinstance(overrides_json, str) else overrides_json
        except (json.JSONDecodeError, TypeError):
            return confirmed

        _flog.info(f"--- _apply_overrides | before={len(confirmed)} entities ---")
        _flog.info(f"  Overrides: remove={overrides.get('remove', [])}, add={len(overrides.get('add', []))} items")

        # Remove false positives: by index AND all matching text+type
        if "remove" in overrides and overrides["remove"]:
            remove_set = set(overrides["remove"])
            removed_signatures = set()
            for i, e in enumerate(confirmed):
                if i in remove_set:
                    removed_signatures.add((e["type"], e["text"].strip().lower()))
                    _flog.info(f"  REMOVE by index {i}: {e['type']} \"{e['text']}\"")
            before_remove = len(confirmed)
            confirmed = [e for i, e in enumerate(confirmed)
                         if i not in remove_set
                         and (e["type"], e["text"].strip().lower()) not in removed_signatures]
            _flog.info(f"  Removed: {before_remove} → {len(confirmed)} entities")

        # Add user-specified entities: find ALL occurrences in text
        if "add" in overrides and overrides["add"]:
            for addition in overrides["add"]:
                add_text = addition.get("text", "")
                add_type = addition.get("type", "PERSON")
                if not add_text:
                    continue
                _flog.info(f"  ADD: type={add_type} text=\"{add_text}\" (searching all occurrences)")
                add_count = 0
                search_start = 0
                while True:
                    pos = text.find(add_text, search_start)
                    if pos < 0:
                        break
                    already_covered = any(
                        e["start"] <= pos and pos + len(add_text) <= e["end"]
                        for e in confirmed
                    )
                    if not already_covered:
                        confirmed.append({
                            "type": add_type,
                            "text": add_text,
                            "start": pos,
                            "end": pos + len(add_text),
                            "score": 1.0,
                            "verified": True,
                            "reason": "user_added",
                        })
                        add_count += 1
                        _flog.info(f"    Added at [{pos}:{pos+len(add_text)}]")
                    else:
                        _flog.info(f"    Skipped at [{pos}:{pos+len(add_text)}] — already covered")
                    search_start = pos + len(add_text)
                _flog.info(f"    Total added for \"{add_text}\": {add_count} occurrences")

        result = sorted(confirmed, key=lambda x: x["start"])
        _flog.info(f"--- _apply_overrides done | after={len(result)} entities ---")
        for e in result:
            _flog.info(f"    {e['type']} [{e['start']}:{e['end']}] \"{e['text']}\" (reason={e.get('reason','NER')})")
        return result

    # --- Text anonymization ---
    def anonymize_text(self, text, language="en", prefix="", entity_overrides=""):
        t0 = time.time()
        _ensure_file_log()
        _flog.info(f"=== anonymize_text() | text_len={len(text)} | prefix={prefix!r} | has_overrides={bool(entity_overrides)} ===")

        entities = self.detect(text, language)
        confirmed = [e for e in entities if e.get("verified")]

        # Apply HITL overrides if provided
        if entity_overrides:
            confirmed = self._apply_overrides(confirmed, text, entity_overrides)

        mapping = self._assign_placeholders(confirmed, prefix)
        _flog.info(f"  Mapping ({len(mapping)} placeholders):")
        for ph, real in sorted(mapping.items()):
            _flog.info(f"    {ph} → \"{real}\"")

        anonymized = text
        for e in sorted(confirmed, key=lambda x: x["start"], reverse=True):
            anonymized = anonymized[:e["start"]] + e["placeholder"] + anonymized[e["end"]:]

        session_id = uuid.uuid4().hex[:12]
        save_mapping(session_id, mapping, {"confirmed": len(confirmed)})

        by_type = defaultdict(int)
        for e in confirmed:
            by_type[e["type"]] += 1

        # Build safe entity list — strip real text, keep only placeholders and metadata
        safe_entities = []
        for e in confirmed:
            safe_entities.append({
                "placeholder": e.get("placeholder", ""),
                "type": e["type"],
                "score": e["score"],
                "verified": e["verified"],
                "reason": e.get("reason", ""),
            })

        # Store review data for potential HITL review (memory + disk for cross-process access)
        review_data = {
            "original_text": text,
            "entities": [{"type": e["type"], "text": e["text"], "start": e["start"],
                          "end": e["end"], "score": e["score"], "verified": e.get("verified", False)}
                         for e in entities],
            "confirmed": [i for i, e in enumerate(entities) if e.get("verified")],
            "status": "pending",
            "overrides": {"remove": [], "add": []},
            "timestamp": time.time(),
        }
        _in_memory_mappings[f"review:{session_id}"] = review_data
        _save_review_to_disk(session_id, review_data)

        return {
            "anonymized_text": anonymized, "session_id": session_id,
            "total_entities": len(entities), "entities_confirmed": len(confirmed),
            "unique_entities": len(mapping),
            "by_type": dict(by_type), "entities": safe_entities,
            "processing_time_ms": round((time.time() - t0) * 1000, 1),
        }

    # --- DOCX anonymization (reuses detect) ---
    def anonymize_docx(self, docx_path, language="en", prefix=""):
        from docx import Document
        t0 = time.time()
        doc = Document(str(docx_path))

        # Shared state across paragraphs (exact dedup + family grouping)
        type_counters = defaultdict(int)
        seen_exact = {}   # (type, exact_normalized_text) → placeholder
        seen_family = {}  # (type, normalized_text) → (family_number, variant_counter)
        mapping = {}      # placeholder → exact raw text
        total = 0
        by_type = defaultdict(int)

        for para in self._iter_docx_paragraphs(doc):
            full_text, runs_info = self._get_runs(para)
            if not full_text.strip():
                continue

            entities = self.detect(full_text, language)
            confirmed = [e for e in entities if e.get("verified")]

            for e in sorted(confirmed, key=lambda x: x["start"], reverse=True):
                ph = self._get_or_create_placeholder(
                    e["type"], e["text"], type_counters, seen_exact, seen_family, mapping, prefix
                )
                self._replace_in_runs(runs_info, e["start"], e["end"], ph)
                total += 1
                by_type[e["type"]] += 1

        docx_path = Path(docx_path)
        session_id = uuid.uuid4().hex[:12]
        out_dir = docx_path.parent / f"hacienda_shield_{session_id}"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{docx_path.stem}_anonymized.docx"
        self._save_docx(doc, out_path)

        save_mapping(session_id, mapping, {"source": str(docx_path)})

        return {
            "output_path": str(out_path), "session_id": session_id,
            "total_entities": total,
            "unique_entities": len(mapping), "by_type": dict(by_type),
            "processing_time_ms": round((time.time() - t0) * 1000, 1),
        }

    # --- Docx save helper (fsync for VirtioFS visibility) ---
    @staticmethod
    def _save_docx(doc, out_path):
        """Save docx and fsync to ensure VirtioFS mount sees the complete file."""
        doc.save(str(out_path))
        try:
            fd = os.open(str(out_path), os.O_RDONLY)
            os.fsync(fd)
            os.close(fd)
        except OSError:
            pass  # fsync not supported on all platforms

    # --- Apply existing mapping to docx (for re-anonymization with HITL overrides) ---
    @staticmethod
    def _collect_paragraph_segments(p_elem, _wns):
        """Collect all inline text-producing elements in document order.
        Returns list of (element, contributed_text, kind) where kind is 'wt'|'br'|'tab'|'cr'.
        w:br → '\\n', w:tab → '\\t', w:cr → '\\r' — matching what para.text produces.
        Skips page/column breaks (w:br with w:type='page'|'column') — only text wrapping breaks."""
        segments = []
        for elem in p_elem.iter():
            tag = elem.tag
            if tag == f'{{{_wns}}}t':
                segments.append((elem, elem.text or "", "wt"))
            elif tag == f'{{{_wns}}}br':
                br_type = elem.get(f'{{{_wns}}}type')
                if br_type in ('page', 'column'):
                    continue  # not a text line break
                segments.append((elem, "\n", "br"))
            elif tag == f'{{{_wns}}}tab':
                segments.append((elem, "\t", "tab"))
            elif tag == f'{{{_wns}}}cr':
                segments.append((elem, "\r", "cr"))
        return segments

    @staticmethod
    def _replace_across_runs(p_elem, old_text, new_text, _wns):
        """Replace old_text with new_text across split runs, preserving formatting.
        Handles w:br (\\n), w:tab (\\t), w:cr (\\r) elements — so text extracted via
        para.text (which includes these chars) can be matched in the XML structure."""
        if not old_text:
            return
        while True:
            segments = PIIEngine._collect_paragraph_segments(p_elem, _wns)
            if not segments:
                break

            joined = "".join(s[1] for s in segments)
            idx = joined.find(old_text)
            if idx == -1:
                break

            end_idx = idx + len(old_text)
            seg_pos = 0
            first_seg = last_seg = -1
            offset_in_first = offset_in_last_end = 0

            for i, (elem, text, kind) in enumerate(segments):
                seg_end = seg_pos + len(text)
                if first_seg == -1 and seg_end > idx:
                    first_seg = i
                    offset_in_first = idx - seg_pos
                if seg_end >= end_idx:
                    last_seg = i
                    offset_in_last_end = end_idx - seg_pos
                    break
                seg_pos = seg_end

            if first_seg == -1 or last_seg == -1:
                break

            # Find first w:t element in match range to host the replacement text
            host_seg = None
            for i in range(first_seg, last_seg + 1):
                if segments[i][2] == "wt":
                    host_seg = i
                    break

            if host_seg is None:
                break  # No text element in range — cannot place replacement

            # Apply replacement across all segments in the match range
            for i in range(first_seg, last_seg + 1):
                elem, text, kind = segments[i]

                if i == host_seg:
                    # This w:t gets the replacement text
                    prefix = text[:offset_in_first] if i == first_seg else ""
                    suffix = text[offset_in_last_end:] if i == last_seg else ""
                    elem.text = prefix + new_text + suffix
                elif kind == "wt":
                    # Other w:t elements: keep only text outside the match
                    if i == first_seg:
                        elem.text = text[:offset_in_first]
                    elif i == last_seg:
                        elem.text = text[offset_in_last_end:]
                    else:
                        elem.text = ""
                else:
                    # Non-text element (br/tab/cr) inside match range: remove from XML
                    parent = elem.getparent()
                    if parent is not None:
                        parent.remove(elem)

    @staticmethod
    def _replace_cross_paragraphs(all_p_elems, old_text, new_text, _wns):
        """Replace text that spans multiple paragraphs (contains \\n from paragraph join).
        Splits old_text by \\n, finds matching consecutive paragraphs, puts replacement in
        first paragraph and clears matched portions in subsequent paragraphs.
        Loops to handle repeated occurrences of the same text."""
        # Filter out empty parts from leading/trailing/consecutive \n
        raw_parts = old_text.split("\n")
        parts = [p for p in raw_parts if p]
        if len(parts) < 2:
            # After filtering empties, if < 2 non-empty parts, can't do cross-paragraph match
            return False

        replaced_any = False
        while True:
            # Rebuild per-paragraph virtual text each iteration (previous replacements change it)
            para_data = []
            for p_elem in all_p_elems:
                segs = PIIEngine._collect_paragraph_segments(p_elem, _wns)
                vtext = "".join(s[1] for s in segs)
                para_data.append((p_elem, vtext))

            found = False
            for start in range(len(para_data) - len(parts) + 1):
                matched = True
                for j, part in enumerate(parts):
                    p_text = para_data[start + j][1]
                    if j == 0:
                        # First part must be a suffix of the paragraph
                        if not part or not p_text.endswith(part):
                            matched = False; break
                    elif j == len(parts) - 1:
                        # Last part must be a prefix of the paragraph
                        if not part or not p_text.startswith(part):
                            matched = False; break
                    else:
                        # Middle parts must match entire paragraph
                        if p_text != part:
                            matched = False; break

                if not matched:
                    continue

                # Match found — apply replacement
                # First paragraph: replace parts[0] at the end with new_text
                PIIEngine._replace_across_runs(all_p_elems[start], parts[0], new_text, _wns)

                # Middle paragraphs: clear all text
                for j in range(1, len(parts) - 1):
                    for seg_elem, _, seg_kind in PIIEngine._collect_paragraph_segments(all_p_elems[start + j], _wns):
                        if seg_kind == "wt":
                            seg_elem.text = ""
                        else:
                            parent = seg_elem.getparent()
                            if parent is not None:
                                parent.remove(seg_elem)

                # Last paragraph: remove parts[-1] from start
                PIIEngine._replace_across_runs(all_p_elems[start + len(parts) - 1], parts[-1], "", _wns)

                found = True
                replaced_any = True
                break  # restart search from beginning (paragraph texts changed)

            if not found:
                break

        return replaced_any

    def anonymize_docx_with_mapping(self, docx_path, mapping, out_dir=None):
        """Apply a known placeholder mapping to a .docx via find-replace. No NER detection.
        Handles split runs, w:br/w:tab/w:cr, and text spanning multiple paragraphs.
        Searches ALL w:p elements in the entire document XML (paragraphs, tables, text boxes,
        content controls, headers, footers, footnotes, endnotes)."""
        from docx import Document
        doc = Document(str(docx_path))
        reverse_map = {v: k for k, v in mapping.items()}
        sorted_texts = sorted(reverse_map.keys(), key=len, reverse=True)
        _wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

        all_p_elems = list(self._iter_all_wp_elements(doc))
        _flog.info(f"=== anonymize_docx_with_mapping | {len(mapping)} placeholders | {len(all_p_elems)} paragraphs ===")
        _flog.info(f"  Reverse map ({len(reverse_map)} entries, sorted by length desc):")
        for rt in sorted_texts:
            _flog.info(f"    \"{rt}\" → {reverse_map[rt]}")

        # Pass 1: single-paragraph replacements (handles split runs + w:br/w:tab/w:cr)
        cross_para_texts = []
        for real_text in sorted_texts:
            found = False
            for p_idx, p_elem in enumerate(all_p_elems):
                segs = self._collect_paragraph_segments(p_elem, _wns)
                vtext = "".join(s[1] for s in segs)
                if real_text in vtext:
                    _flog.info(f"  PASS1 FOUND: \"{real_text[:60]}\" in para {p_idx} → {reverse_map[real_text]}")
                    self._replace_across_runs(p_elem, real_text, reverse_map[real_text], _wns)
                    found = True
            if not found and "\n" in real_text:
                cross_para_texts.append(real_text)
                _flog.info(f"  PASS1 NOT FOUND (cross-para): \"{real_text[:60]}\" → queued for pass 2")
            elif not found:
                _flog.warning(f"  PASS1 NOT FOUND: \"{real_text[:60]}\" → {reverse_map[real_text]} — WILL NOT BE ANONYMIZED IN DOCX")

        # Pass 2: cross-paragraph replacements for texts spanning multiple paragraphs
        if cross_para_texts:
            _flog.info(f"  PASS2: {len(cross_para_texts)} cross-paragraph texts to replace")
            for real_text in cross_para_texts:
                all_p_fresh = list(self._iter_all_wp_elements(doc))
                result = self._replace_cross_paragraphs(all_p_fresh, real_text, reverse_map[real_text], _wns)
                if result:
                    _flog.info(f"  PASS2 OK: \"{real_text[:60]}\" → {reverse_map[real_text]}")
                else:
                    _flog.warning(f"  PASS2 FAILED: \"{real_text[:60]}\" → {reverse_map[real_text]} — NOT ANONYMIZED")

        _flog.info(f"=== anonymize_docx_with_mapping done ===")
        parent = Path(out_dir) if out_dir else Path(docx_path).parent
        out = parent / f"{Path(docx_path).stem}_anonymized.docx"
        self._save_docx(doc, out)
        return str(out)

    # --- Deanonymization ---
    @staticmethod
    def deanonymize_text(text, mapping):
        for ph in sorted(mapping.keys(), key=len, reverse=True):
            text = text.replace(ph, mapping[ph])
        return text

    def deanonymize_docx(self, docx_path, mapping):
        """Restore placeholders in .docx — handles split runs, w:br/w:tab/w:cr, and cross-paragraph.
        Searches ALL w:p elements in the entire document XML tree."""
        from docx import Document
        doc = Document(str(docx_path))
        sorted_ph = sorted(mapping.keys(), key=len, reverse=True)
        _wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        # Placeholders like <PERSON_1> don't contain \n, so single-pass is enough
        for p_elem in self._iter_all_wp_elements(doc):
            for ph in sorted_ph:
                self._replace_across_runs(p_elem, ph, mapping[ph], _wns)
        out = Path(docx_path).parent / f"{Path(docx_path).stem}_restored.docx"
        self._save_docx(doc, out)
        return str(out)

    # --- DOCX helpers ---
    @staticmethod
    def _iter_docx_paragraphs(doc):
        """Iterate paragraphs via python-docx API (body, tables, headers/footers).
        Used for text extraction where para.text / para.style are needed."""
        for p in doc.paragraphs:
            yield p
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for sec in doc.sections:
            for hf in [sec.header, sec.footer]:
                if hf:
                    for p in hf.paragraphs:
                        yield p

    @staticmethod
    def _is_inside_tracked_delete(p_elem, _wns):
        """Check if a w:p element is inside a w:del (tracked deletion) ancestor."""
        parent = p_elem.getparent()
        while parent is not None:
            if parent.tag == f'{{{_wns}}}del':
                return True
            parent = parent.getparent()
        return False

    @staticmethod
    def _iter_all_wp_elements(doc):
        """Iterate ALL w:p elements in the entire document XML tree.
        Catches: body paragraphs, tables, text boxes (w:txbxContent),
        content controls (w:sdtContent), tracked insertions, headers, footers.
        Skips w:p elements inside w:del (tracked deletions) to avoid corrupting revision history."""
        _wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        # Main document body
        for p in doc.element.iter(f'{{{_wns}}}p'):
            if not PIIEngine._is_inside_tracked_delete(p, _wns):
                yield p
        # Headers, footers (they have separate XML parts)
        for sec in doc.sections:
            for part in [sec.header, sec.footer, getattr(sec, 'first_page_header', None),
                         getattr(sec, 'first_page_footer', None),
                         getattr(sec, 'even_page_header', None),
                         getattr(sec, 'even_page_footer', None)]:
                if part and part._element is not None:
                    for p in part._element.iter(f'{{{_wns}}}p'):
                        if not PIIEngine._is_inside_tracked_delete(p, _wns):
                            yield p

    @staticmethod
    def _docx_to_html(doc):
        """Convert docx to simple HTML preserving bold/italic/underline and headings.
        Text content matches para.text exactly: includes hyperlinks, excludes
        tracked changes (w:ins/w:del), handles w:br as newlines."""
        from html import escape as _html_escape
        from lxml import etree
        _nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        _wns = _nsmap['w']

        def _run_formatting(r_elem):
            """Extract bold/italic/underline from w:rPr."""
            rPr = r_elem.find(f'{{{_wns}}}rPr')
            if rPr is None:
                return False, False, False
            b = rPr.find(f'{{{_wns}}}b')
            bold = b is not None and b.get(f'{{{_wns}}}val', 'true') != 'false'
            i = rPr.find(f'{{{_wns}}}i')
            italic = i is not None and i.get(f'{{{_wns}}}val', 'true') != 'false'
            u = rPr.find(f'{{{_wns}}}u')
            underline = u is not None and u.get(f'{{{_wns}}}val', 'none') != 'none'
            return bold, italic, underline

        parts = []
        for para in PIIEngine._iter_docx_paragraphs(doc):
            style_name = para.style.name if para.style else ""
            tag = "p"
            if "Heading 1" in style_name or "Title" in style_name:
                tag = "h1"
            elif "Heading 2" in style_name or "Subtitle" in style_name:
                tag = "h2"
            elif "Heading 3" in style_name:
                tag = "h3"
            elif "Heading" in style_name:
                tag = "h4"

            # Match para.text exactly: only direct child w:r and w:hyperlink,
            # then their inline children (w:t, w:br, w:tab, w:cr, w:noBreakHyphen).
            # para.text uses xpath 'w:r|w:hyperlink' on direct children only —
            # NOT w:r buried inside w:smartTag, w:sdt, w:fldSimple, w:ins, w:del, etc.
            runs_html = []

            def _process_run(r_elem):
                """Process a single w:r element — extract text with formatting."""
                bold, italic, underline = _run_formatting(r_elem)
                for child in r_elem:
                    child_tag = etree.QName(child.tag).localname if '}' in child.tag else child.tag
                    if child_tag == 't':
                        if child.text:
                            t = _html_escape(child.text)
                            if bold:
                                t = f"<b>{t}</b>"
                            if italic:
                                t = f"<i>{t}</i>"
                            if underline:
                                t = f"<u>{t}</u>"
                            runs_html.append(t)
                    elif child_tag == 'br':
                        # Only textWrapping (or no type) = \n. Page/column breaks = empty.
                        br_type = child.get(f'{{{_wns}}}type')
                        if br_type is None or br_type == 'textWrapping':
                            runs_html.append('<br>')
                    elif child_tag in ('tab', 'ptab'):
                        runs_html.append('&#9;')
                    elif child_tag == 'cr':
                        runs_html.append('<br>')
                    elif child_tag == 'noBreakHyphen':
                        runs_html.append('-')

            for child in para._element:
                child_tag = etree.QName(child.tag).localname if '}' in child.tag else child.tag
                if child_tag == 'r':
                    _process_run(child)
                elif child_tag == 'hyperlink':
                    for sub in child:
                        sub_tag = etree.QName(sub.tag).localname if '}' in sub.tag else sub.tag
                        if sub_tag == 'r':
                            _process_run(sub)
            parts.append(f"<{tag}>{''.join(runs_html)}</{tag}>")
        return "\n".join(parts)

    @staticmethod
    def _get_runs(para):
        runs_info = []
        offset = 0
        for run in para.runs:
            runs_info.append({"run": run, "text": run.text, "start": offset, "end": offset + len(run.text)})
            offset += len(run.text)
        return "".join(r["text"] for r in runs_info), runs_info

    @staticmethod
    def _replace_in_runs(runs_info, start, end, replacement):
        """Replace text at [start, end) with replacement across runs.

        Uses a two-phase approach: modify run texts first, then recalculate
        all offsets from scratch. This avoids cascading offset errors when
        an entity spans multiple runs.
        """
        # Phase 1: identify affected runs and modify their text
        affected = []
        for i, ri in enumerate(runs_info):
            if ri["end"] <= start or ri["start"] >= end:
                continue
            affected.append(i)

        if not affected:
            return

        first_idx = affected[0]
        last_idx = affected[-1]

        for idx in affected:
            ri = runs_info[idx]
            old_text = ri["text"]
            local_start = max(0, start - ri["start"])
            local_end = min(len(old_text), end - ri["start"])

            if idx == first_idx and idx == last_idx:
                # Entity fully within one run
                new_text = old_text[:local_start] + replacement + old_text[local_end:]
            elif idx == first_idx:
                # Entity starts here, continues into next run(s)
                new_text = old_text[:local_start] + replacement
            elif idx == last_idx:
                # Entity ends here, started in previous run(s)
                new_text = old_text[local_end:]
            else:
                # Middle run — entirely within entity range
                new_text = ""

            ri["run"].text = new_text
            ri["text"] = new_text

        # Phase 2: recalculate all offsets from scratch
        offset = 0
        for ri in runs_info:
            ri["start"] = offset
            ri["end"] = offset + len(ri["text"])
            offset += len(ri["text"])


# ============================================================
# MCP Tools
# ============================================================
mcp = FastMCP("Hacienda Shield", host="127.0.0.1", port=int(os.environ.get("PII_PORT", "8765")), lifespan=_pii_lifespan)
engine = PIIEngine()

# --- Chunked processing state ---
_chunk_sessions = {}  # session_id -> chunk state dict
_CHUNK_SESSION_TTL = 1800  # 30 minutes


def _split_paragraphs(text, target_size):
    """Split text into chunks on paragraph boundaries (\\n\\n).
    Each chunk is at most target_size chars (unless a single paragraph exceeds it)."""
    paragraphs = text.split("\n\n")
    chunks = []
    current = []
    current_len = 0
    for para in paragraphs:
        para_len = len(para) + (2 if current else 0)  # +2 for \n\n separator
        if current_len + para_len > target_size and current:
            chunks.append("\n\n".join(current))
            current = [para]
            current_len = len(para)
        else:
            current.append(para)
            current_len += para_len
    if current:
        chunks.append("\n\n".join(current))
    return chunks if chunks else [text]


def _process_chunk(session_id):
    """Process the current chunk: detect entities, assign placeholders with shared state."""
    cs = _chunk_sessions[session_id]
    chunk_idx = cs["current_chunk"]
    chunk_text = cs["chunks"][chunk_idx]

    # Detect entities in this chunk
    engine._ensure_ready()
    entities = engine.detect(chunk_text, cs["language"])
    confirmed = [e for e in entities if e.get("verified")]

    # Calculate offset for this chunk in the full text
    offset = 0
    for i in range(chunk_idx):
        offset += len(cs["chunks"][i]) + 2  # +2 for \n\n separator

    # Assign placeholders using shared state across all chunks
    for e in sorted(confirmed, key=lambda x: x["start"]):
        e["start"] += offset
        e["end"] += offset
        e["placeholder"] = engine._get_or_create_placeholder(
            e["type"], e["text"],
            cs["type_counters"], cs["seen_exact"], cs["seen_family"],
            cs["mapping"], cs["prefix"]
        )

    cs["all_entities"].extend(confirmed)
    cs["current_chunk"] = chunk_idx + 1
    return confirmed


def _cleanup_stale_chunk_sessions():
    """Remove chunk sessions older than TTL."""
    now = time.time()
    stale = [sid for sid, cs in _chunk_sessions.items()
             if now - cs.get("created_at", 0) > _CHUNK_SESSION_TTL]
    for sid in stale:
        del _chunk_sessions[sid]


def _latest_session_id():
    """Find the most recent session by mapping file mtime, with in-memory fallback."""
    try:
        sessions = sorted(
            (f for f in MAPPING_DIR.glob("*.json") if not f.name.startswith("review_")),
            key=lambda p: p.stat().st_mtime, reverse=True
        )
        if sessions:
            return json.loads(sessions[0].read_text(encoding="utf-8")).get("session_id", "")
    except Exception:
        pass
    # Fallback to in-memory
    if _in_memory_mappings:
        latest = max(_in_memory_mappings.values(), key=lambda d: d.get("timestamp", 0))
        return latest.get("session_id", "")
    return ""


def _check_ready():
    """Check if engine is ready. Returns None if ready, or a JSON status string if still loading."""
    if _engine_ready.is_set() and not _boot_error:
        return None  # Ready
    if _boot_error:
        return json.dumps({
            "status": "error",
            "message": f"Hacienda Shield failed to initialize: {_boot_error}",
            "hint": "Check internet connection, ensure Python 3.10+ and restart.",
        }, indent=2)
    # Still loading — use benchmark for ETA
    elapsed = round(time.monotonic() - _boot_progress["start"], 1) if _boot_progress.get("start") else 0
    bench = _load_boot_benchmark()
    eta_total = bench.get("total_sec")
    eta_remaining = max(0, round(eta_total - elapsed, 1)) if eta_total else None
    return json.dumps({
        "status": "loading",
        "phase": _boot_progress.get("phase", "starting"),
        "message": _boot_progress.get("message", "Hacienda Shield is starting up..."),
        "progress_pct": _boot_progress.get("pct", 0),
        "elapsed_seconds": elapsed,
        "eta_remaining_sec": eta_remaining,
        "retry_after_sec": 10 if elapsed > 180 else 25,
    }, indent=2)


@mcp.tool()
@_audit_tool
def anonymize_text(text: str, language: str = "en", prefix: str = "", entity_overrides: str = "") -> str:
    """Anonymize PII in text. Returns indexed placeholders + session_id for deanonymization.
    Use prefix (e.g. "D1") for multi-file workflows to avoid placeholder collisions.
    Use entity_overrides (JSON) from HITL review to add/remove entities."""
    loading = _check_ready()
    if loading:
        return loading
    r = engine.anonymize_text(text, language, prefix=prefix, entity_overrides=entity_overrides)
    return json.dumps(r, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def anonymize_file(file_path: str, language: str = "en", prefix: str = "", review_session_id: str = "") -> str:
    """Anonymize PII in a file. Auto-detects format: .pdf, .docx (preserves formatting), .txt/.md/.csv (plain text).
    Use prefix (e.g. "D1") for multi-file workflows to avoid placeholder collisions.
    Use review_session_id to re-anonymize with HITL overrides — the server fetches overrides internally, PII never passes through the API.
    PREFERRED over anonymize_text for privacy: only the file path passes through the API, not the content."""
    loading = _check_ready()
    if loading:
        return loading
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        # Try work_dir + filename
        work_dir = os.environ.get("PII_WORK_DIR", "").strip()
        if work_dir:
            candidate = Path(work_dir).expanduser().resolve() / p.name
            if candidate.exists():
                p = candidate
        if not p.exists():
            return json.dumps({"error": f"File not found: {p}",
                               "hint": "Ask the user for the full host path to the file."})

    # Create output subdirectory to keep generated files organized
    # Re-anonymization reuses the same output dir via session_id
    def _make_output_dir(parent, sid):
        out_dir = parent / f"hacienda_shield_{sid}"
        out_dir.mkdir(parents=True, exist_ok=True)
        return out_dir

    # Resolve HITL overrides from review session (stored on server, never sent to Claude)
    entity_overrides = ""
    if review_session_id:
        review = _get_review(review_session_id.strip())
        if review:
            overrides = review.get("overrides", {})
            if overrides.get("remove") or overrides.get("add"):
                entity_overrides = json.dumps(overrides)
        else:
            return json.dumps({"error": f"Review session not found: {review_session_id}. Run anonymize_file + start_review first."})

    # --- Extract text based on format ---
    text = None
    docx_html = None

    if p.suffix.lower() == ".pdf":
        try:
            import pdfplumber
        except ImportError:
            try:
                import subprocess
                subprocess.check_call([sys.executable, "-m", "pip", "install", "pdfplumber", "-q"],
                                      stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                import pdfplumber
            except Exception as e:
                return json.dumps({"error": f"Cannot install pdfplumber for PDF support: {e}"})
        try:
            with pdfplumber.open(str(p)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            return json.dumps({"error": f"Failed to read PDF: {e}"})
        if len(text.strip()) < 50:
            return json.dumps({"error": "PDF has no extractable text layer. Scanned PDFs (OCR) are not yet supported.",
                               "file": str(p)})
    elif p.suffix.lower() == ".docx":
        _ensure_file_log(str(p.parent))
        _flog.info(f"======== anonymize_file DOCX | {p.name} | review_session_id={review_session_id!r} ========")
        try:
            from docx import Document as _DocxDoc
            _doc = _DocxDoc(str(p))
            # Body paragraphs as lines
            parts = [para.text for para in _doc.paragraphs]
            # Table rows: join cells with " | " so NER gets context for short values
            # e.g. "James Whitfield | Managing Director | 15 March 2016"
            for t in _doc.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            # Headers/footers
            for sec in _doc.sections:
                for hf in [sec.header, sec.footer]:
                    if hf:
                        for para in hf.paragraphs:
                            parts.append(para.text)
            text = "\n".join(parts)
            _flog.info(f"  Extracted text: {len(text)} chars from {len(parts)} parts")
        except Exception as e:
            return json.dumps({"error": f"Failed to read docx: {e}"})
        try:
            docx_html = PIIEngine._docx_to_html(_doc)
        except Exception:
            docx_html = None
    elif p.suffix.lower() in (".txt", ".md", ".csv", ".log", ".text"):
        text = p.read_text(encoding="utf-8")
    else:
        return json.dumps({"error": f"Unsupported format: {p.suffix}. Supported: .pdf .docx .txt .md .csv"})

    # --- Chunked processing for long documents (>15K chars) ---
    if len(text) > 15000:
        _cleanup_stale_chunk_sessions()
        engine._ensure_ready()

        # Auto-calibrate: measure NER speed on a small sample
        TIMEOUT_BUDGET = 40  # seconds, with margin before 60s MCP timeout
        cal_sample = text[:3000]
        t_cal = time.time()
        engine.detect(cal_sample, language)
        cal_elapsed = time.time() - t_cal
        chars_per_sec = 3000 / max(cal_elapsed, 0.1)
        optimal_chunk_size = int(chars_per_sec * TIMEOUT_BUDGET)
        optimal_chunk_size = max(4000, min(50000, optimal_chunk_size))

        chunks = _split_paragraphs(text, optimal_chunk_size)
        if len(chunks) > 1:
            chunk_session_id = uuid.uuid4().hex[:12]
            _chunk_sessions[chunk_session_id] = {
                "text": text,
                "chunks": chunks,
                "current_chunk": 0,
                "type_counters": defaultdict(int),
                "seen_exact": {},
                "seen_family": {},
                "mapping": {},
                "all_entities": [],
                "prefix": prefix,
                "language": language,
                "source_path": str(p),
                "source_suffix": p.suffix.lower(),
                "chars_per_sec": chars_per_sec,
                "optimal_chunk_size": optimal_chunk_size,
                "entity_overrides": entity_overrides,
                "docx_html": docx_html,
                "created_at": time.time(),
            }

            # Process first chunk immediately
            first_confirmed = _process_chunk(chunk_session_id)

            cs = _chunk_sessions[chunk_session_id]
            return json.dumps({
                "status": "chunked",
                "session_id": chunk_session_id,
                "total_chunks": len(chunks),
                "processed_chunks": 1,
                "progress_pct": round(100 / len(chunks)),
                "entities_so_far": len(first_confirmed),
                "chars_per_sec": round(chars_per_sec),
                "chunk_size": optimal_chunk_size,
                "note": "Document is large. Call anonymize_next_chunk(session_id) to process remaining chunks, then get_full_anonymized_text(session_id).",
            }, indent=2, ensure_ascii=False)

    # --- Standard (non-chunked) processing ---
    r = engine.anonymize_text(text, language, prefix=prefix, entity_overrides=entity_overrides)
    out_dir = _make_output_dir(p.parent, r["session_id"])

    if p.suffix.lower() == ".docx":
        _ensure_file_log(str(out_dir))
        _flog.info(f"  Output dir: {out_dir}")
        if docx_html:
            review_key = f"review:{r['session_id']}"
            if review_key in _in_memory_mappings:
                _in_memory_mappings[review_key]["original_html"] = docx_html
                _save_review_to_disk(r['session_id'], _in_memory_mappings[review_key])
        out_txt = out_dir / f"{p.stem}_anonymized.txt"
        out_txt.write_text(r["anonymized_text"], encoding="utf-8")
        r.pop("anonymized_text", None)
        r["output_path"] = str(out_txt)
        r["output_dir"] = str(out_dir)
        try:
            mapping = load_mapping(r["session_id"])
            _flog.info(f"  Producing anonymized .docx with {len(mapping)} placeholders")
            docx_out = engine.anonymize_docx_with_mapping(p, mapping, out_dir)
            r["docx_output_path"] = docx_out
            _flog.info(f"  Docx saved: {docx_out}")
        except Exception as e:
            log.warning(f"anonymize_docx failed (txt output OK): {e}")
            _flog.warning(f"  anonymize_docx FAILED: {e}")
        r["note"] = "Anonymized text at output_path (.txt). For REDLINE, use docx_output_path (.docx with formatting)."
    elif p.suffix.lower() == ".pdf":
        out = out_dir / f"{p.stem}_anonymized.txt"
        out.write_text(r["anonymized_text"], encoding="utf-8")
        r.pop("anonymized_text", None)
        r["output_path"] = str(out)
        r["output_dir"] = str(out_dir)
        r["note"] = "Anonymized text written to output_path. Read the file to get the content."
    else:
        out = out_dir / f"{p.stem}_anonymized{p.suffix}"
        out.write_text(r["anonymized_text"], encoding="utf-8")
        r.pop("anonymized_text", None)
        r["output_path"] = str(out)
        r["output_dir"] = str(out_dir)
        r["note"] = "Anonymized text written to output_path. Read the file to get the content."

    return json.dumps(r, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def anonymize_next_chunk(session_id: str) -> str:
    """Process next chunk of a chunked anonymization session.
    Call repeatedly after anonymize_file returns status='chunked'.
    Returns progress (processed_chunks, total_chunks, progress_pct)."""
    if session_id not in _chunk_sessions:
        return json.dumps({"error": f"Chunk session not found: {session_id}",
                           "hint": "Session may have expired (30 min TTL) or was already finalized."})

    cs = _chunk_sessions[session_id]
    if cs["current_chunk"] >= len(cs["chunks"]):
        return json.dumps({"status": "complete",
                           "session_id": session_id,
                           "total_entities": len(cs["all_entities"]),
                           "note": "All chunks processed. Call get_full_anonymized_text(session_id) to finalize."})

    # Process next chunk with timing
    t0 = time.time()
    confirmed = _process_chunk(session_id)
    elapsed = time.time() - t0
    chunk_len = len(cs["chunks"][cs["current_chunk"] - 1])

    # Adaptive: if speed dropped >30%, re-split remaining chunks smaller
    if elapsed > 0.1:
        new_speed = chunk_len / elapsed
        if new_speed < cs["chars_per_sec"] * 0.7:
            cs["chars_per_sec"] = (cs["chars_per_sec"] + new_speed) / 2  # sliding average
            new_size = max(4000, min(50000, int(cs["chars_per_sec"] * 40)))
            # Re-split only remaining unprocessed text
            if cs["current_chunk"] < len(cs["chunks"]):
                remaining_chunks = cs["chunks"][cs["current_chunk"]:]
                remaining_text = "\n\n".join(remaining_chunks)
                new_chunks = _split_paragraphs(remaining_text, new_size)
                cs["chunks"] = cs["chunks"][:cs["current_chunk"]] + new_chunks
                log.info(f"Chunk session {session_id}: speed drop, re-split remaining into {len(new_chunks)} chunks (was {len(remaining_chunks)})")

    is_complete = cs["current_chunk"] >= len(cs["chunks"])
    return json.dumps({
        "status": "complete" if is_complete else "in_progress",
        "session_id": session_id,
        "processed_chunks": cs["current_chunk"],
        "total_chunks": len(cs["chunks"]),
        "progress_pct": round(100 * cs["current_chunk"] / len(cs["chunks"])),
        "entities_so_far": len(cs["all_entities"]),
        "note": "All chunks processed. Call get_full_anonymized_text(session_id) to finalize." if is_complete
                else "Call anonymize_next_chunk(session_id) again for the next chunk.",
    }, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def get_full_anonymized_text(session_id: str) -> str:
    """Assemble all processed chunks, finalize mapping, write output file.
    Call after all chunks are processed (anonymize_next_chunk returned status='complete')."""
    if session_id not in _chunk_sessions:
        return json.dumps({"error": f"Chunk session not found: {session_id}",
                           "hint": "Session may have expired or was already finalized."})

    cs = _chunk_sessions[session_id]

    # Process any remaining chunks
    while cs["current_chunk"] < len(cs["chunks"]):
        _process_chunk(session_id)

    # Apply entity_overrides (from HITL review) on the full entity list
    if cs.get("entity_overrides"):
        cs["all_entities"] = engine._apply_overrides(
            cs["all_entities"], cs["text"], cs["entity_overrides"]
        )
        # Re-assign placeholders for any new entities added by overrides
        for e in cs["all_entities"]:
            if "placeholder" not in e:
                e["placeholder"] = engine._get_or_create_placeholder(
                    e["type"], e["text"],
                    cs["type_counters"], cs["seen_exact"], cs["seen_family"],
                    cs["mapping"], cs["prefix"]
                )

    # Build anonymized text by replacing entities in reverse order
    anonymized = cs["text"]
    for e in sorted(cs["all_entities"], key=lambda x: x["start"], reverse=True):
        anonymized = anonymized[:e["start"]] + e["placeholder"] + anonymized[e["end"]:]

    # Save mapping with a final session_id
    final_session_id = uuid.uuid4().hex[:12]
    save_mapping(final_session_id, cs["mapping"], {"confirmed": len(cs["all_entities"])})

    # Write output file
    p = Path(cs["source_path"])
    out_dir = p.parent / f"hacienda_shield_{final_session_id}"
    out_dir.mkdir(parents=True, exist_ok=True)

    by_type = defaultdict(int)
    for e in cs["all_entities"]:
        by_type[e["type"]] += 1

    # Store review data for HITL
    all_entities_raw = [{"type": e["type"], "text": e["text"], "start": e["start"],
                         "end": e["end"], "score": e["score"], "verified": e.get("verified", False)}
                        for e in cs["all_entities"]]
    review_data = {
        "original_text": cs["text"],
        "entities": all_entities_raw,
        "confirmed": list(range(len(all_entities_raw))),
        "status": "pending",
        "overrides": {"remove": [], "add": []},
        "timestamp": time.time(),
    }
    if cs.get("docx_html"):
        review_data["original_html"] = cs["docx_html"]
    _in_memory_mappings[f"review:{final_session_id}"] = review_data
    _save_review_to_disk(final_session_id, review_data)

    result = {
        "session_id": final_session_id,
        "total_entities": len(cs["all_entities"]),
        "unique_entities": len(cs["mapping"]),
        "by_type": dict(by_type),
        "output_dir": str(out_dir),
    }

    if cs["source_suffix"] == ".docx":
        out_txt = out_dir / f"{p.stem}_anonymized.txt"
        out_txt.write_text(anonymized, encoding="utf-8")
        result["output_path"] = str(out_txt)
        try:
            docx_out = engine.anonymize_docx_with_mapping(p, cs["mapping"], out_dir)
            result["docx_output_path"] = docx_out
        except Exception as e:
            log.warning(f"Chunked docx output failed: {e}")
        result["note"] = "Anonymized text at output_path (.txt). For REDLINE, use docx_output_path (.docx with formatting)."
    elif cs["source_suffix"] == ".pdf":
        out_txt = out_dir / f"{p.stem}_anonymized.txt"
        out_txt.write_text(anonymized, encoding="utf-8")
        result["output_path"] = str(out_txt)
        result["note"] = "Anonymized text written to output_path."
    else:
        out_file = out_dir / f"{p.stem}_anonymized{cs['source_suffix']}"
        out_file.write_text(anonymized, encoding="utf-8")
        result["output_path"] = str(out_file)
        result["note"] = "Anonymized text written to output_path."

    # Cleanup chunk session
    del _chunk_sessions[session_id]

    return json.dumps(result, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def find_file(filename: str) -> str:
    """Find a file on the host machine by filename. Searches the configured work_dir (Settings > Extensions > Hacienda Shield).
    If work_dir is not set or file not found there, returns an error — ask the user for the path."""
    work_dir = os.environ.get("PII_WORK_DIR", "").strip()
    if not work_dir:
        return json.dumps({"error": "Working directory not configured.",
                           "hint": "Ask the user for the full file path, or ask them to set 'Working directory' in Settings > Extensions > Hacienda Shield."})
    wd = Path(work_dir).expanduser().resolve()
    if not wd.exists():
        return json.dumps({"error": f"Configured work_dir does not exist: {work_dir}",
                           "hint": "Ask the user to fix 'Working directory' in Settings > Extensions > Hacienda Shield."})
    matches = []
    try:
        for f in wd.rglob(filename):
            if f.is_file():
                matches.append(str(f))
                if len(matches) >= 10:
                    break
    except PermissionError:
        pass
    if matches:
        return json.dumps({"matches": matches, "count": len(matches)})
    return json.dumps({"error": f"File '{filename}' not found in work_dir: {work_dir}",
                       "hint": "Ask the user for the full file path."})


# --- Marker-based path resolution (VM → Host) ---

_BFS_SKIP = {".git", "node_modules", "__pycache__", ".venv", "venv", ".tox",
             ".mypy_cache", "$Recycle.Bin", "System Volume Information", "Windows"}
_dir_cache = {}  # vm_dir -> host_dir
_DIR_CACHE_FILE = Path.home() / ".hacienda_shield" / "dir_cache.json"


def _load_dir_cache():
    """Load cached VM→Host directory mappings from disk."""
    global _dir_cache
    if _DIR_CACHE_FILE.exists():
        try:
            _dir_cache = json.loads(_DIR_CACHE_FILE.read_text(encoding="utf-8"))
        except Exception:
            _dir_cache = {}


def _save_dir_cache():
    """Save VM→Host directory mappings to disk."""
    try:
        _DIR_CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
        _DIR_CACHE_FILE.write_text(json.dumps(_dir_cache, indent=2), encoding="utf-8")
    except Exception:
        pass


def _bfs_find(root, target_name, max_depth=6):
    """Depth-limited BFS for a file by exact name. Returns Path or None."""
    from collections import deque
    queue = deque([(Path(root), 0)])
    while queue:
        current, depth = queue.popleft()
        try:
            for entry in current.iterdir():
                if entry.name == target_name and entry.is_file():
                    return entry
                if (entry.is_dir()
                        and depth < max_depth
                        and not entry.name.startswith('.')
                        and entry.name not in _BFS_SKIP):
                    queue.append((entry, depth + 1))
        except (PermissionError, OSError):
            continue
    return None


def _find_marker(marker, max_depth=6):
    """Search for a unique marker file across home dir and drive roots."""
    # Home directory first (most likely location)
    result = _bfs_find(Path.home(), marker, max_depth)
    if result:
        return result

    # Platform-specific additional roots
    if sys.platform == "win32":
        import string
        home_drive = Path.home().drive
        for letter in string.ascii_uppercase:
            drive = Path(f"{letter}:\\")
            if f"{letter}:" == home_drive:
                continue  # already searched via home
            if drive.exists():
                result = _bfs_find(drive, marker, max_depth)
                if result:
                    return result
    elif sys.platform == "darwin":
        volumes = Path("/Volumes")
        if volumes.exists():
            for vol in volumes.iterdir():
                if vol.is_dir() and vol.name != "Macintosh HD":
                    result = _bfs_find(vol, marker, max_depth)
                    if result:
                        return result
    else:  # Linux
        for root in [Path("/home"), Path("/mnt"), Path("/media")]:
            if root.exists():
                result = _bfs_find(root, marker, max_depth)
                if result:
                    return result
    return None


# Load dir cache on module import
_load_dir_cache()


@mcp.tool()
@_audit_tool
def resolve_path(filename: str, marker: str, vm_dir: str = "") -> str:
    """Zero-config file path resolution between VM and host.
    Claude creates a unique marker file next to the target file in the VM.
    This tool finds that marker on the host via BFS, derives the host path.

    Args:
        filename: Target file name (e.g. 'contract.docx')
        marker: Exact marker filename Claude created (e.g. '.pii_marker_a1b2c3d4')
        vm_dir: Optional VM directory path for caching the mapping
    """
    # Check cache first
    cache_key = vm_dir or marker
    if cache_key in _dir_cache:
        cached_dir = Path(_dir_cache[cache_key])
        candidate = cached_dir / filename
        if candidate.exists():
            # Clean up marker if it exists
            marker_path = cached_dir / marker
            if marker_path.exists():
                try:
                    marker_path.unlink()
                except Exception:
                    pass
            return json.dumps({"host_path": str(candidate), "host_dir": str(cached_dir), "cached": True})

    # BFS search for marker
    found = _find_marker(marker)
    if not found:
        return json.dumps({
            "error": f"Marker file '{marker}' not found on host.",
            "hint": "Ensure the marker was created in the connected folder. Ask the user for the full host path as fallback.",
        })

    host_dir = found.parent

    # Clean up marker
    try:
        found.unlink()
    except Exception:
        pass

    # Cache the mapping
    _dir_cache[cache_key] = str(host_dir)
    if vm_dir:
        _dir_cache[vm_dir] = str(host_dir)
    _save_dir_cache()

    # Find the target file
    target = host_dir / filename
    if target.exists():
        return json.dumps({"host_path": str(target), "host_dir": str(host_dir)})

    return json.dumps({
        "error": f"Marker found at {host_dir} but '{filename}' not in that directory.",
        "host_dir": str(host_dir),
        "hint": "Check the filename spelling.",
    })


@mcp.tool()
@_audit_tool
def anonymize_docx(file_path: str, language: str = "en", prefix: str = "") -> str:
    """Anonymize PII in .docx preserving all formatting. Use for round-trip document editing.
    Use prefix (e.g. "D1") for multi-file workflows to avoid placeholder collisions."""
    loading = _check_ready()
    if loading:
        return loading
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        return json.dumps({"error": f"Not found: {p}"})
    r = engine.anonymize_docx(p, language, prefix=prefix)
    return json.dumps(r, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def deanonymize_text(text: str, session_id: str = "", output_path: str = "") -> str:
    """Restore real PII values in text. Writes result to .docx file — never returns PII to Claude.
    Returns only the file path. output_path should end with .docx (default) or .txt."""
    sid = session_id.strip() or _latest_session_id()
    if not sid:
        return json.dumps({"error": "No session. Run anonymize first."})
    mapping = load_mapping(sid)
    if not mapping:
        return json.dumps({"error": f"Mapping not found: {sid}"})
    restored = engine.deanonymize_text(text, mapping)
    # Determine output path
    if output_path:
        out_p = Path(output_path).expanduser().resolve()
    else:
        out_p = MAPPING_DIR / f"restored_{sid}.docx"
    out_p.parent.mkdir(parents=True, exist_ok=True)
    # Write to file — PII never goes back to Claude
    if out_p.suffix.lower() == ".docx":
        _write_docx(restored, out_p)
    else:
        out_p.write_text(restored, encoding="utf-8")
    return json.dumps({
        "restored_path": str(out_p),
        "session_id": sid,
        "entities_restored": len(mapping),
        "note": "PII-safe: restored text written to file, not returned to LLM.",
    }, indent=2, ensure_ascii=False)


def _write_docx(text: str, path: Path):
    """Write text to a formatted .docx file."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        # Detect headings by common patterns
        if stripped.isupper() and len(stripped) < 100:
            p = doc.add_paragraph(stripped)
            p.style = doc.styles["Heading 2"]
        elif stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            p = doc.add_paragraph(stripped.lstrip("# "))
            p.style = doc.styles[f"Heading {level}"]
        else:
            doc.add_paragraph(stripped)
    PIIEngine._save_docx(doc, path)


@mcp.tool()
@_audit_tool
def deanonymize_docx(file_path: str, session_id: str = "") -> str:
    """Restore real PII in .docx preserving formatting."""
    sid = session_id.strip() or _latest_session_id()
    if not sid:
        return json.dumps({"error": "No session. Run anonymize first."})
    mapping = load_mapping(sid)
    if not mapping:
        return json.dumps({"error": f"Mapping not found: {sid}"})
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        return json.dumps({"error": f"Not found: {p}"})
    out = engine.deanonymize_docx(p, mapping)
    return json.dumps({"restored_path": out, "session_id": sid}, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def get_mapping(session_id: str = "") -> str:
    """Retrieve mapping metadata (placeholder keys and entity types only — no real PII values).
    Full mapping stays on disk, never returned to LLM."""
    sid = session_id.strip() or _latest_session_id()
    if not sid:
        return json.dumps({"error": "No session available."})
    mapping = load_mapping(sid)
    if not mapping:
        return json.dumps({"error": f"Not found: {sid}"})
    # Return only placeholder keys and types — no real PII values
    safe_summary = {}
    for placeholder in mapping:
        # Extract type from placeholder like <PERSON_1> → PERSON
        etype = placeholder.strip("<>").rsplit("_", 1)[0] if "_" in placeholder else placeholder.strip("<>")
        safe_summary[placeholder] = etype
    return json.dumps({
        "session_id": sid,
        "total_entities": len(mapping),
        "placeholders": safe_summary,
        "note": "PII-safe: real values not returned. Use deanonymize_text/docx to restore to file.",
    }, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def scan_text(text: str, language: str = "en") -> str:
    """Detect PII without anonymizing. Preview mode. Returns entity types and positions, not real text values."""
    loading = _check_ready()
    if loading:
        return loading
    entities = engine.detect(text, language)
    # Strip real text from entities — return only type, position, and verification status
    safe_entities = []
    for e in entities:
        safe_entities.append({
            "type": e["type"],
            "start": e["start"],
            "end": e["end"],
            "score": e["score"],
            "verified": e.get("verified"),
            "reason": e.get("reason", ""),
        })
    return json.dumps({
        "entities_found": len(entities),
        "confirmed": len([e for e in entities if e.get("verified")]),
        "rejected": len([e for e in entities if not e.get("verified")]),
        "entities": safe_entities,
    }, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def list_entities() -> str:
    """Show status, supported types, and recent sessions."""
    # Always show recent sessions (no engine needed)
    sessions = sorted(
        (f for f in MAPPING_DIR.glob("*.json") if not f.name.startswith("review_")),
        key=lambda p: p.stat().st_mtime, reverse=True
    )[:5]
    recent = []
    for s in sessions:
        try:
            d = json.loads(s.read_text())
            recent.append({"session_id": d["session_id"], "entities": len(d["mapping"])})
        except Exception:
            pass

    # If still loading, show status without crashing
    if not _engine_ready.is_set():
        return json.dumps({
            "status": "loading",
            "phase": _boot_progress.get("phase", "starting"),
            "message": _boot_progress.get("message", "Hacienda Shield is starting up..."),
            "progress_pct": _boot_progress.get("pct", 0),
            "retry_after_sec": 10 if (_boot_progress.get("start") and (_time_boot.monotonic() - _boot_progress["start"]) > 180) else 25,
            "recent_sessions": recent,
        }, indent=2, ensure_ascii=False)

    if _boot_error:
        return json.dumps({
            "status": "error",
            "message": f"Model load failed: {_boot_error}",
            "recent_sessions": recent,
        }, indent=2, ensure_ascii=False)

    # Engine ready — full diagnostics
    eng = PIIEngine()
    eng._ensure_ready()
    recognizer_names = [type(r).__name__ for r in eng.analyzer.registry.recognizers]
    engine_class = type(eng.analyzer.nlp_engine).__name__

    backend = eng._backend or ""
    if "gliner" in backend:
        quality = "full (GLiNER zero-shot)"
    elif "FALLBACK" not in backend:
        quality = "full"
    else:
        quality = "reduced (SpaCy fallback)"
    return json.dumps({
        "status": "ready",
        "mode": "NER-only",
        "backend": eng._backend,
        "quality": quality,
        "nlp_engine_class": engine_class,
        "recognizers": recognizer_names,
        "min_score": _get_min_score(),
        "mapping_ttl_days": MAPPING_TTL_DAYS,
        "recent_sessions": recent,
    }, indent=2, ensure_ascii=False)


# ============================================================
# HITL Review Web Server (localhost only)
# ============================================================
import threading
from http.server import HTTPServer, BaseHTTPRequestHandler

_review_server = None
_review_port = None
# Search multiple possible locations for review_ui.html
_REVIEW_HTML_PATH = None
for _candidate in [
    Path(__file__).parent / "review_ui.html",                          # same dir as script
    Path(__file__).parent.parent / "server" / "review_ui.html",        # if script moved up one level
    Path(__file__).parent.parent / "review_ui.html",                   # parent dir
    Path(os.environ.get("PII_SHIELD_DIR", "")) / "server" / "review_ui.html",  # env override
]:
    if _candidate.exists():
        _REVIEW_HTML_PATH = _candidate
        break
if _REVIEW_HTML_PATH:
    _blog.info(f"Review UI found at: {_REVIEW_HTML_PATH}")
else:
    _blog.warning(f"review_ui.html NOT FOUND. Searched near: {Path(__file__).parent}")
    _REVIEW_HTML_PATH = Path(__file__).parent / "review_ui.html"  # fallback for error message


class _ReviewHandler(BaseHTTPRequestHandler):
    """Localhost-only review UI handler. PII never leaves the machine."""

    def log_message(self, fmt, *args):
        """Suppress default HTTP logging to stderr (interferes with MCP stdio)."""
        pass

    def _send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _send_html(self, html):
        body = html.encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_GET(self):
        path = self.path.split("?")[0]
        if path.startswith("/review/"):
            session_id = path.split("/review/")[1]
            self._serve_review_page(session_id)
        elif path.startswith("/api/review/"):
            session_id = path.split("/api/review/")[1]
            self._serve_review_data(session_id)
        else:
            self.send_error(404)

    def do_POST(self):
        path = self.path.split("?")[0]
        content_length = int(self.headers.get("Content-Length", 0))
        body = self.rfile.read(content_length) if content_length > 0 else b""

        if "/api/approve/" in path:
            session_id = path.split("/api/approve/")[1]
            self._handle_approve(session_id, body)
        elif "/api/remove_entity/" in path:
            session_id = path.split("/api/remove_entity/")[1]
            self._handle_remove_entity(session_id, body)
        elif "/api/add_entity/" in path:
            session_id = path.split("/api/add_entity/")[1]
            self._handle_add_entity(session_id, body)
        else:
            self.send_error(404)

    def _serve_review_page(self, session_id):
        if not _get_review(session_id):
            self._send_html(f"<h1>Review session not found: {session_id}</h1>")
            return
        try:
            html = _REVIEW_HTML_PATH.read_text(encoding="utf-8")
            self._send_html(html)
        except FileNotFoundError:
            self._send_html(f"<h1>review_ui.html not found</h1><p>Searched: {_REVIEW_HTML_PATH}</p><p>Script: {Path(__file__).resolve()}</p>")

    def _serve_review_data(self, session_id):
        review = _get_review(session_id)
        if not review:
            self._send_json({"error": f"Review session not found: {session_id}"}, 404)
            return
        data = {
            "session_id": session_id,
            "original_text": review["original_text"],
            "entities": review["entities"],
            "confirmed": review["confirmed"],
            "status": review["status"],
            "overrides": review["overrides"],
        }
        if "original_html" in review:
            data["original_html"] = review["original_html"]
        self._send_json(data)

    def _handle_approve(self, session_id, body):
        review = _load_review_from_disk(session_id) or _get_review(session_id)
        if not review:
            self._send_json({"error": "Session not found"}, 404)
            return
        try:
            overrides = json.loads(body) if body else {}
        except json.JSONDecodeError:
            overrides = {}
        review["status"] = "approved"
        review["overrides"] = {
            "remove": overrides.get("remove", []),
            "add": overrides.get("add", []),
        }
        _save_review_to_disk(session_id, review)
        self._send_json({"status": "approved", "session_id": session_id})

    def _handle_remove_entity(self, session_id, body):
        # Reload from disk to avoid stale data from other processes
        review = _load_review_from_disk(session_id) or _get_review(session_id)
        if not review:
            self._send_json({"error": "Session not found"}, 404)
            return
        try:
            data = json.loads(body)
            idx = data.get("index")
        except (json.JSONDecodeError, TypeError):
            self._send_json({"error": "Invalid JSON"}, 400)
            return
        if idx is not None and isinstance(idx, int) and 0 <= idx < len(review.get("entities", [])):
            if idx not in review["overrides"]["remove"]:
                review["overrides"]["remove"].append(idx)
                _save_review_to_disk(session_id, review)
        self._send_json({"ok": True, "overrides": review["overrides"]})

    def _handle_add_entity(self, session_id, body):
        # Reload from disk to avoid stale data from other processes
        review = _load_review_from_disk(session_id) or _get_review(session_id)
        if not review:
            self._send_json({"error": "Session not found"}, 404)
            return
        try:
            data = json.loads(body)
        except (json.JSONDecodeError, TypeError):
            self._send_json({"error": "Invalid JSON"}, 400)
            return
        text = data.get("text", "").strip()
        start = data.get("start", -1)
        end = data.get("end", -1)
        if not text or start < 0 or end <= start:
            self._send_json({"error": "Invalid entity: need text, start >= 0, end > start"}, 400)
            return
        review["overrides"]["add"].append({
            "text": text,
            "type": data.get("type", "PERSON"),
            "start": start,
            "end": end,
        })
        _save_review_to_disk(session_id, review)
        self._send_json({"ok": True, "overrides": review["overrides"]})


def _start_review_server():
    """Start localhost-only review web server in a background thread."""
    global _review_server, _review_port
    if _review_server is not None:
        return _review_port
    _review_port = int(os.environ.get("PII_REVIEW_PORT", "8766"))
    try:
        server = HTTPServer(("127.0.0.1", _review_port), _ReviewHandler)
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        _review_server = server
        log.info(f"Review server started on http://127.0.0.1:{_review_port}")
    except OSError as e:
        log.warning(f"Could not start review server on port {_review_port}: {e}")
        # Try next port
        _review_port += 1
        try:
            server = HTTPServer(("127.0.0.1", _review_port), _ReviewHandler)
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            _review_server = server
            log.info(f"Review server started on http://127.0.0.1:{_review_port}")
        except OSError:
            log.error("Failed to start review server")
            return None
    return _review_port


@mcp.tool()
@_audit_tool
def start_review(session_id: str = "") -> str:
    """Start local review server and return the URL. Does NOT open the browser — Claude presents the link to the user via AskUserQuestion.
    PII stays on your machine."""
    sid = session_id.strip() or _latest_session_id()
    if not sid:
        return json.dumps({"error": "No session. Run anonymize_file first."})
    review = _get_review(sid)
    if not review:
        return json.dumps({"error": f"No review data for session {sid}. Run anonymize_file first."})
    port = _start_review_server()
    if port is None:
        return json.dumps({"error": "Could not start review server."})
    log.info(f"Review HTML path: {_REVIEW_HTML_PATH} (exists: {_REVIEW_HTML_PATH.exists()})")
    log.info(f"Server script at: {Path(__file__).resolve()}")
    url = f"http://localhost:{port}/review/{sid}"
    entity_count = len([i for i in review.get("confirmed", [])])
    return json.dumps({
        "url": url,
        "session_id": sid,
        "entities_count": entity_count,
        "note": "Review server ready. Present this URL to the user via AskUserQuestion. Do NOT open browser automatically.",
    }, indent=2, ensure_ascii=False)


@mcp.tool()
@_audit_tool
def get_review_status(session_id: str = "") -> str:
    """Check if user approved the HITL review. Returns status and whether changes were made.
    PII-safe: never returns override details (entity text). Use review_session_id in anonymize_file to apply them."""
    sid = session_id.strip() or _latest_session_id()
    if not sid:
        return json.dumps({"error": "No session available."})
    review = _get_review(sid)
    if not review:
        return json.dumps({"error": f"No review for session {sid}"})
    overrides = review.get("overrides", {"remove": [], "add": []})
    has_changes = bool(overrides.get("remove") or overrides.get("add"))
    return json.dumps({
        "session_id": sid,
        "status": review["status"],
        "has_changes": has_changes,
        "removed_count": len(overrides.get("remove", [])),
        "added_count": len(overrides.get("add", [])),
    }, indent=2, ensure_ascii=False)


def _ensure_ssl_cert(cert_dir: Path):
    """Generate self-signed cert if not exists."""
    cert_file = cert_dir / "cert.pem"
    key_file = cert_dir / "key.pem"
    if cert_file.exists() and key_file.exists():
        return str(cert_file), str(key_file)
    cert_dir.mkdir(parents=True, exist_ok=True)
    log.info("Generating self-signed SSL certificate...")
    import subprocess
    subprocess.run([
        sys.executable, "-c",
        f"""
from cryptography import x509
from cryptography.x509.oid import NameOID
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import rsa
import datetime, ipaddress
key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
name = x509.Name([x509.NameAttribute(NameOID.COMMON_NAME, "localhost")])
cert = (x509.CertificateBuilder()
    .subject_name(name).issuer_name(name)
    .public_key(key.public_key())
    .serial_number(x509.random_serial_number())
    .not_valid_before(datetime.datetime.now(datetime.timezone.utc))
    .not_valid_after(datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(days=3650))
    .add_extension(x509.SubjectAlternativeName([
        x509.DNSName("localhost"),
        x509.IPAddress(ipaddress.IPv4Address("127.0.0.1")),
    ]), critical=False)
    .sign(key, hashes.SHA256()))
open(r"{cert_file}", "wb").write(cert.public_bytes(serialization.Encoding.PEM))
open(r"{key_file}", "wb").write(key.private_bytes(
    serialization.Encoding.PEM, serialization.PrivateFormat.TraditionalOpenSSL, serialization.NoEncryption()))
print("OK")
"""
    ], check=True)
    log.info(f"SSL cert created: {cert_file}")
    return str(cert_file), str(key_file)


if __name__ == "__main__":
    transport = "sse" if "--sse" in sys.argv else "stdio"
    port = int(os.environ.get("PII_PORT", "8765"))

    log.info(f"Starting Hacienda Shield MCP Server v1.0.0 ({transport})...")

    if transport == "sse":
            import ssl
            import uvicorn

            cert_dir = Path.home() / ".hacienda-shield" / "ssl"
            cert_file, key_file = _ensure_ssl_cert(cert_dir)

            app = mcp.sse_app()
            uvicorn.run(
                app,
                host="127.0.0.1",
                port=port,
                ssl_certfile=cert_file,
                ssl_keyfile=key_file,
                log_level="info",
            )
    else:
        mcp.run(transport="stdio")
